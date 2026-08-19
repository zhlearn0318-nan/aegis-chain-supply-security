from __future__ import annotations

import argparse
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


NAVY = "17365D"
BLUE = "2E74B5"
MID_BLUE = "5B9BD5"
LIGHT_BLUE = "DCE6F1"
PALE_BLUE = "EEF4FA"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "6B7280"
DARK = "1F2937"
GREEN = "2F855A"
PALE_GREEN = "E9F5EE"
GOLD = "A66A00"
PALE_GOLD = "FFF4D6"
RED = "A61B1B"
PALE_RED = "FCE8E6"
WHITE = "FFFFFF"

ASCII_FONT = "Calibri"
EAST_ASIA_BODY = "SimSun"
EAST_ASIA_HEAD = "Microsoft YaHei"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_run_font(
    run,
    *,
    ascii_name: str = ASCII_FONT,
    east_asia_name: str = EAST_ASIA_BODY,
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
):
    run.font.name = ascii_name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), ascii_name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), ascii_name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia_name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = TABLE_INDENT_DXA) -> None:
    if sum(widths_dxa) != TABLE_WIDTH_DXA:
        raise ValueError(f"Table widths must sum to {TABLE_WIDTH_DXA}, got {sum(widths_dxa)}")

    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent_dxa))

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[min(index, len(widths_dxa) - 1)]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color="D8DEE8", size=6) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_paragraph_keep(paragraph, keep_next=True, keep_lines=True) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    if keep_next:
        keep = p_pr.find(qn("w:keepNext"))
        if keep is None:
            keep = OxmlElement("w:keepNext")
            p_pr.append(keep)
    if keep_lines:
        keep_lines_el = p_pr.find(qn("w:keepLines"))
        if keep_lines_el is None:
            keep_lines_el = OxmlElement("w:keepLines")
            p_pr.append(keep_lines_el)


def add_field(paragraph, field_code: str, display: str = "1") -> None:
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_char)
    instr_run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {field_code} "
    instr_run._r.append(instr)
    sep_run = paragraph.add_run()
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    sep_run._r.append(sep)
    value_run = paragraph.add_run(display)
    set_run_font(value_run, size=9, color=MID_GRAY)
    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def configure_styles(doc: Document) -> None:
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = ASCII_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), ASCII_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), ASCII_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_BODY)
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(DARK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    normal_p_pr = normal._element.get_or_add_pPr()
    if normal_p_pr.find(qn("w:keepLines")) is None:
        normal_p_pr.append(OxmlElement("w:keepLines"))

    title = styles["Title"]
    title.font.name = ASCII_FONT
    title._element.rPr.rFonts.set(qn("w:ascii"), ASCII_FONT)
    title._element.rPr.rFonts.set(qn("w:hAnsi"), ASCII_FONT)
    title._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_HEAD)
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = rgb(NAVY)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)
    title.paragraph_format.line_spacing = 1.05

    subtitle = styles["Subtitle"]
    subtitle.font.name = ASCII_FONT
    subtitle._element.rPr.rFonts.set(qn("w:ascii"), ASCII_FONT)
    subtitle._element.rPr.rFonts.set(qn("w:hAnsi"), ASCII_FONT)
    subtitle._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_HEAD)
    subtitle.font.size = Pt(12.5)
    subtitle.font.color.rgb = rgb(MID_GRAY)
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(14)

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, NAVY, 8, 4),
    ):
        style = styles[style_name]
        style.font.name = ASCII_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), ASCII_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), ASCII_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_HEAD)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = ASCII_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), ASCII_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), ASCII_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_BODY)
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    caption = styles["Caption"]
    caption.font.name = ASCII_FONT
    caption._element.rPr.rFonts.set(qn("w:ascii"), ASCII_FONT)
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), ASCII_FONT)
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_BODY)
    caption.font.size = Pt(9)
    caption.font.color.rgb = rgb(MID_GRAY)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(4)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


def configure_page(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)


def configure_header_footer(doc: Document) -> None:
    for section in doc.sections:
        section.different_first_page_header_footer = True
        header = section.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run("XA-202620  |  供应链安全模块阶段性汇报")
        set_run_font(r, east_asia_name=EAST_ASIA_HEAD, size=8.5, color=MID_GRAY, bold=True)

        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.paragraph_format.space_before = Pt(0)
        fp.paragraph_format.space_after = Pt(0)
        r = fp.add_run("第 ")
        set_run_font(r, size=9, color=MID_GRAY)
        add_field(fp, "PAGE")
        r = fp.add_run(" 页 / 共 ")
        set_run_font(r, size=9, color=MID_GRAY)
        add_field(fp, "NUMPAGES")
        r = fp.add_run(" 页")
        set_run_font(r, size=9, color=MID_GRAY)

        first_footer = section.first_page_footer
        ffp = first_footer.paragraphs[0]
        ffp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = ffp.add_run("挑战杯揭榜挂帅赛道 · 阶段性技术汇报")
        set_run_font(r, east_asia_name=EAST_ASIA_HEAD, size=8.5, color=MID_GRAY)


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    set_paragraph_keep(p)
    return p


def add_body(doc: Document, text: str, *, bold_prefix: str | None = None, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, bold=True)
        r = p.add_run(text[len(bold_prefix):])
        set_run_font(r)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    set_run_font(r)
    return p


def create_decimal_numbering(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
        if node.get(qn("w:abstractNumId")) is not None
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
        if node.get(qn("w:numId")) is not None
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract_num = OxmlElement("w:abstractNum")
    abstract_num.set(qn("w:abstractNumId"), str(abstract_id))
    multi_level = OxmlElement("w:multiLevelType")
    multi_level.set(qn("w:val"), "singleLevel")
    abstract_num.append(multi_level)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%1.")
    level_jc = OxmlElement("w:lvlJc")
    level_jc.set(qn("w:val"), "left")
    level.extend([start, num_fmt, level_text, level_jc])
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.extend([tabs, ind, spacing])
    level.append(p_pr)
    abstract_num.append(level)
    first_num_index = next(
        (index for index, child in enumerate(numbering) if child.tag == qn("w:num")),
        len(numbering),
    )
    numbering.insert(first_num_index, abstract_num)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_number(doc: Document, text: str, num_id: int, *, keep_next: bool = False):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])
    p_pr.append(num_pr)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    if keep_next:
        set_paragraph_keep(p, keep_next=True, keep_lines=True)
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_callout(doc: Document, label: str, text: str, *, fill=PALE_BLUE, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    set_table_borders(table, color=accent, size=8)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.10
    r = p.add_run(label + "  ")
    set_run_font(r, east_asia_name=EAST_ASIA_HEAD, size=11, color=accent, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=11, color=DARK)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    return table


def add_table_caption(doc: Document, text: str):
    p = doc.add_paragraph(text, style="Caption")
    set_paragraph_keep(p)
    return p


def add_data_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    widths_dxa: list[int],
    *,
    alignments: list[int] | None = None,
    header_fill: str = LIGHT_GRAY,
    font_size: float = 9.5,
):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_dxa)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(header)
        set_run_font(r, east_asia_name=EAST_ASIA_HEAD, size=font_size, color=NAVY, bold=True)
    for row_values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_values):
            cell = cells[i]
            if len(table.rows) % 2 == 0:
                set_cell_shading(cell, "FBFCFE")
            p = cell.paragraphs[0]
            p.alignment = (alignments[i] if alignments else WD_ALIGN_PARAGRAPH.LEFT)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.05
            r = p.add_run(str(value))
            set_run_font(r, size=font_size, color=DARK)
    set_table_geometry(table, widths_dxa)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), ASCII_FONT)
    r_fonts.set(qn("w:hAnsi"), ASCII_FONT)
    r_fonts.set(qn("w:eastAsia"), EAST_ASIA_BODY)
    r_pr.extend([r_fonts, color, underline])
    run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def create_charts(asset_dir: Path) -> tuple[Path, Path]:
    asset_dir.mkdir(parents=True, exist_ok=True)
    font_regular = "C:/Windows/Fonts/arial.ttf"
    font_bold = "C:/Windows/Fonts/arialbd.ttf"
    label_font = ImageFont.truetype(font_regular, 29)
    value_font = ImageFont.truetype(font_bold, 27)
    tick_font = ImageFont.truetype(font_regular, 24)
    axis_font = ImageFont.truetype(font_regular, 25)

    def draw_vertical_grid(draw, left, top, right, bottom, *, maximum=100, step=20):
        for value in range(0, maximum + 1, step):
            x = left + (right - left) * value / maximum
            draw.line((x, top, x, bottom), fill="#D9DEE7", width=2)
            label = str(value)
            box = draw.textbbox((0, 0), label, font=tick_font)
            draw.text((x - (box[2] - box[0]) / 2, bottom + 12), label, font=tick_font, fill="#5B6472")

    core_path = asset_dir / "core_metrics.png"
    labels = ["Coverage", "Loose precision", "Non-normal recall", "Malicious recall", "Normal FPR"]
    values = [98.20, 87.94, 76.67, 72.28, 24.70]
    colors = [f"#{MID_BLUE}", f"#{GREEN}", f"#{BLUE}", f"#{NAVY}", f"#{RED}"]
    image = Image.new("RGB", (1584, 720), "white")
    draw = ImageDraw.Draw(image)
    left, top, right, bottom = 390, 40, 1450, 620
    draw_vertical_grid(draw, left, top, right, bottom)
    row_height = (bottom - top) / len(labels)
    bar_height = 58
    for index, (label, value, color) in enumerate(zip(labels, values, colors)):
        cy = top + row_height * (index + 0.5)
        label_box = draw.textbbox((0, 0), label, font=label_font)
        draw.text((left - 24 - (label_box[2] - label_box[0]), cy - 18), label, font=label_font, fill="#303946")
        x2 = left + (right - left) * value / 100
        draw.rounded_rectangle((left, cy - bar_height / 2, x2, cy + bar_height / 2), radius=10, fill=color)
        value_text = f"{value:.2f}%"
        value_box = draw.textbbox((0, 0), value_text, font=value_font)
        tx = min(x2 + 16, right - (value_box[2] - value_box[0]))
        if tx < x2 + 8:
            tx = x2 - (value_box[2] - value_box[0]) - 12
            value_color = "white"
        else:
            value_color = "#1F2937"
        draw.text((tx, cy - 17), value_text, font=value_font, fill=value_color)
    axis_text = "Percent (%)"
    axis_box = draw.textbbox((0, 0), axis_text, font=axis_font)
    draw.text(((left + right) / 2 - (axis_box[2] - axis_box[0]) / 2, 674), axis_text, font=axis_font, fill="#5B6472")
    image.save(core_path, dpi=(220, 220))

    risk_path = asset_dir / "risk_recall.png"
    risk_labels = [f"T{i:02d}" for i in range(1, 10)]
    risk_values = [82.31, 75.00, 76.92, 79.24, 66.67, 40.00, 91.67, 96.00, 62.50]
    risk_colors = [f"#{BLUE}" if v >= 70 else f"#{GOLD}" if v >= 60 else f"#{RED}" for v in risk_values]
    image = Image.new("RGB", (1584, 690), "white")
    draw = ImageDraw.Draw(image)
    left, top, right, bottom = 135, 50, 1500, 580
    for value in range(0, 101, 20):
        y = bottom - (bottom - top) * value / 100
        draw.line((left, y, right, y), fill="#D9DEE7", width=2)
        label = str(value)
        box = draw.textbbox((0, 0), label, font=tick_font)
        draw.text((left - 18 - (box[2] - box[0]), y - 14), label, font=tick_font, fill="#5B6472")
    slot = (right - left) / len(risk_labels)
    bar_width = slot * 0.60
    for index, (label, value, color) in enumerate(zip(risk_labels, risk_values, risk_colors)):
        cx = left + slot * (index + 0.5)
        y = bottom - (bottom - top) * value / 100
        draw.rounded_rectangle((cx - bar_width / 2, y, cx + bar_width / 2, bottom), radius=7, fill=color)
        value_text = f"{value:.1f}"
        box = draw.textbbox((0, 0), value_text, font=value_font)
        draw.text((cx - (box[2] - box[0]) / 2, y - 38), value_text, font=value_font, fill="#1F2937")
        box = draw.textbbox((0, 0), label, font=label_font)
        draw.text((cx - (box[2] - box[0]) / 2, bottom + 18), label, font=label_font, fill="#303946")
    draw.text((left, 8), "Recall (%)", font=axis_font, fill="#5B6472")
    image.save(risk_path, dpi=(220, 220))
    return core_path, risk_path


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run("挑战杯揭榜挂帅赛道 · 阶段性技术汇报")
    set_run_font(r, east_asia_name=EAST_ASIA_HEAD, size=10.5, color=BLUE, bold=True)

    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run("XA-202620 面向政企场景的\n大模型智能体安全关键技术研究")
    p.runs[0].font.name = ASCII_FONT
    p.runs[0]._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), EAST_ASIA_HEAD)

    p = doc.add_paragraph(style="Subtitle")
    r = p.add_run("供应链安全模块静态审查阶段性汇报\n——基于 Cisco Skill Scanner 与 SkillTrustBench 官方 556 条固定子集")
    set_run_font(r, east_asia_name=EAST_ASIA_HEAD, size=12.5, color=MID_GRAY)

    metadata = [
        ("汇报人", "陶泽泓"),
        ("模块定位", "智能体 Skill / MCP 供应链安全"),
        ("当前阶段", "Cisco 静态底座接入与权威公开基准评测"),
        ("报告日期", "2026 年 8 月 14 日"),
        ("实验结论", "accepted_with_caveats（有条件接受）"),
    ]
    table = doc.add_table(rows=len(metadata), cols=2)
    set_table_geometry(table, [1800, 7560])
    set_table_borders(table, color="D9E2F0", size=5)
    for row, (label, value) in zip(table.rows, metadata):
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        set_cell_shading(row.cells[1], "FBFCFE")
        p0 = row.cells[0].paragraphs[0]
        p0.paragraph_format.space_before = Pt(3)
        p0.paragraph_format.space_after = Pt(3)
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r0 = p0.add_run(label)
        set_run_font(r0, east_asia_name=EAST_ASIA_HEAD, size=10, color=NAVY, bold=True)
        p1 = row.cells[1].paragraphs[0]
        p1.paragraph_format.space_before = Pt(3)
        p1.paragraph_format.space_after = Pt(3)
        r1 = p1.add_run(value)
        set_run_font(r1, size=10, color=DARK)
    set_table_geometry(table, [1800, 7560])

    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    add_callout(
        doc,
        "核心判断",
        "系统已形成可运行、可追溯、可批量复现的静态准入门，但当前 Cisco 静态结果不能独立承担最终安全裁决。应以漏报切片补强、上下文降误报和最小动态验证闭环作为下一阶段重点。",
        fill=PALE_BLUE,
        accent=BLUE,
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("内部汇报材料 · 数据与指标均可由逐样本结果复算")
    set_run_font(r, east_asia_name=EAST_ASIA_HEAD, size=9, color=MID_GRAY, italic=True)

    doc.add_page_break()


def build_report(output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    asset_dir = output_path.parent / ".docx_qa_assets"
    core_chart, risk_chart = create_charts(asset_dir)

    doc = Document()
    configure_page(doc)
    configure_styles(doc)
    configure_header_footer(doc)
    doc.core_properties.title = "XA-202620 供应链安全模块静态审查阶段性汇报"
    doc.core_properties.subject = "Cisco Skill Scanner × SkillTrustBench 官方 556 条子集评测"
    doc.core_properties.author = "陶泽泓"
    doc.core_properties.keywords = "XA-202620, 供应链安全, Cisco Skill Scanner, SkillTrustBench, 智能体安全"
    doc.core_properties.comments = "由项目评测证据生成的正式阶段性汇报稿"
    update_fields = doc.settings.element.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        doc.settings.element.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    add_cover(doc)

    add_heading(doc, "摘  要", 1)
    add_body(
        doc,
        "本报告面向 XA-202620 赛题中“供应链安全”方向，汇报 Cisco Skill Scanner 静态审查底座的接入、评测与下一阶段开发计划。项目选用腾讯朱雀实验室与香港中文大学（深圳）联合发布的 SkillTrustBench 公开基准，固定其官方 10% 子集 556 条，在 Windows 11 本地环境中关闭 LLM、云上传与行为执行，仅启用静态分析、字节码分析和流水线文本/结构检查。",
    )
    add_body(
        doc,
        "本轮 556/556 条样本均得到终态，其中 546 条完成 Cisco 扫描，7 条发生运行时错误，3 条真实恶意样本被 Windows Defender 阻断读取；10 条失败均按 UNKNOWN/abstain 处理，没有伪装成安全结果。最终覆盖率 98.20%，严格三分类 macro F1 为 0.4977，恶意严格召回率 72.28%；在“REVIEW/BLOCK 均视为识别到风险”的统一门禁口径下，风险筛查 precision 为 87.94%、recall 为 76.67%、loose F1 为 81.92%，正常样本误报率为 24.70%。",
    )
    add_body(
        doc,
        "结果表明：现有系统已完成从第三方扫描器到统一证据模型、准入策略、批量评测、失败闭锁和可复现追踪的工程闭环，但静态规则对 wild real-world、crypto_wallet 以及 T06/T09/T05 风险仍存在明显漏报；同时，网络请求、文件访问等通用能力规则存在上下文不足导致的误报。下一阶段应优先围绕 91 条风险筛查漏报做规则补强，并用无害模拟样本建立进程、网络、文件和敏感环境变量四类动态事件的最小验证闭环。",
    )

    add_heading(doc, "汇报结论摘要", 2)
    for text in (
        "工程结论：Cisco Skill Scanner 已可作为本系统静态审查底座稳定运行；当前单条中位耗时约 4.0 秒，覆盖率约 98%。",
        "能力结论：风险门禁 loose F1 达 81.92%，说明 REVIEW/BLOCK 的筛查价值已形成；严格三分类 macro F1 仅 0.4977，说明 normal、suspicious、malicious 的精细区分仍不足。",
        "安全结论：所有异常均失败闭锁为 UNKNOWN，未关闭 Windows Defender、未执行样本、未安装依赖、未上传云端，避免了“为了跑分绕过安全控制”。",
        "下一步结论：先补漏报和上下文判断，再补最小动态验证；不把 Cisco 分数直接当作系统创新，也不宣称复现官方 Cisco 榜单成绩。",
    ):
        add_bullet(doc, text)

    add_heading(doc, "1 研究背景与阶段目标", 1)
    add_heading(doc, "1.1 问题背景", 2)
    add_body(
        doc,
        "大模型智能体通过 Skill、MCP Server、插件和外部依赖获得工具能力，也同时引入新的软件供应链风险。攻击者可能把恶意指令、隐蔽网络访问、凭据读取、持久化逻辑或不安全依赖嵌入工具包。对于政企场景，仅依赖模型输出过滤无法覆盖这些风险，需要在工具接入前建立可审计、可追溯的供应链准入机制。",
    )
    add_heading(doc, "1.2 本阶段目标", 2)
    stage_goal_items = (
        "验证 Cisco Skill Scanner 和 MCP Scanner 在本地环境中的可运行性，并明确两者适用对象。",
        "先完成 Skill 侧静态审查主链路：原始扫描输出适配、统一 Finding、策略映射、四态决策、失败闭锁和批量评测。",
        "使用公开、可固定版本的基准数据集形成可复现的外部证据，而不是只用少量自造样本展示。",
        "为后续动态审计和统一平台对接预留标准化结果接口。",
    )
    stage_goal_num_id = create_decimal_numbering(doc)
    for index, text in enumerate(stage_goal_items):
        add_number(doc, text, stage_goal_num_id, keep_next=index < len(stage_goal_items) - 1)

    add_heading(doc, "1.3 模块边界", 2)
    add_callout(
        doc,
        "当前边界",
        "SkillTrustBench 用于评估 Skill 静态扫描能力；Cisco MCP Scanner 面向 MCP server/config，不强行套用到 Skill 数据集。当前报告只评价本地离线静态门禁，不包含真实恶意样本动态执行，也不等同于完整系统最终成绩。",
        fill=PALE_GOLD,
        accent=GOLD,
    )

    add_heading(doc, "2 技术方案与工程实现", 1)
    add_heading(doc, "2.1 静态审查主链路", 2)
    pipeline_items = (
        "安全导入：固定数据集 revision 和文件哈希，审计 ZIP 路径穿越、符号链接、成员数量与解压规模，只解压指定 case。",
        "第三方扫描：以独立子进程调用 Cisco Skill Scanner，仅允许 static_analyzer、bytecode、pipeline 三类本地分析器。",
        "证据归一：将厂商原始结果转换为统一 Finding，只保留规则 ID、类别、严重度、分析器和位置等脱敏信息。",
        "策略决策：通过冻结的 Aegis Chain 策略输出 ALLOW、REVIEW、BLOCK 或 UNKNOWN，并记录策略版本与 SHA-256。",
        "评测追踪：逐 case 保存终态、预测、耗时和树哈希；支持断点续扫、指标复算和失败诊断。",
    )
    pipeline_num_id = create_decimal_numbering(doc)
    for index, text in enumerate(pipeline_items):
        add_number(doc, text, pipeline_num_id, keep_next=index < len(pipeline_items) - 1)

    add_table_caption(doc, "表 1  四态准入决策与评测映射")
    add_data_table(
        doc,
        ["系统决策", "三分类映射", "运营含义", "是否自动放行"],
        [
            ["ALLOW", "normal", "未发现达到门槛的风险证据", "是"],
            ["REVIEW", "suspicious", "存在风险信号，进入人工复核", "否"],
            ["BLOCK", "malicious", "命中高风险规则，阻断接入", "否"],
            ["UNKNOWN", "abstain", "扫描失败或证据不足，按失败闭锁处理", "否"],
        ],
        [1500, 1800, 4260, 1800],
        alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER],
    )

    add_heading(doc, "2.2 安全与可复现设计", 2)
    for text in (
        "不安装 Skill，不导入样本 Python 模块，不执行样本脚本、Shell 命令或依赖。",
        "关闭 LLM、VirusTotal、AI Defense 和其他云分析凭据，避免恶意样本内容外传。",
        "扫描前后复核样本 tree SHA-256；553 条可读样本均未发生内容变化。",
        "保留运行 ID、扫描器版本/哈希、策略版本/哈希、数据清单哈希和逐样本结果，保证结果可追溯。",
        "出现运行时错误或终端防护阻断时，不把样本判为安全，而是进入 UNKNOWN。",
    ):
        add_bullet(doc, text)

    add_heading(doc, "3 数据集与实验设计", 1)
    add_heading(doc, "3.1 数据来源与固定身份", 2)
    add_body(
        doc,
        "SkillTrustBench 是面向智能体 Skill 安全扫描器的公开基准，完整 v1.0 audited refresh 含 5,520 条 normal / suspicious / malicious 三类样本。官方结果仓库提供固定 10% 评测清单，本项目严格使用该 556 条清单，不因扫描结果更换样本。该数据集适合作为竞赛阶段的公开外部基准，但截至审计时未找到专门介绍其构造与标注方法的同行评审数据集论文，因此正式表述为“联合发布的公开基准”，不称为国际标准。",
    )
    add_table_caption(doc, "表 2  官方 556 条固定子集的可复现身份")
    add_data_table(
        doc,
        ["项目", "固定值"],
        [
            ["数据集提交", "762d5388b3a047b26df9679582af868a0e5b2c8f"],
            ["结果仓库提交", "326ec286d082199cb270b25b8b4fc93c8762281e"],
            ["子集文件 SHA-256", "dff7621ffcc7a42f1a8ff64c8e47d2fafc1cd332431fd533be88bb684aaa6843"],
            ["排序 ID 清单 SHA-256", "903a036e4b7b16ee28e22d5d9db57a00b3764cfe41e43144acad67921e5196c2"],
            ["样本与标签", "556：normal 166 / suspicious 105 / malicious 285"],
            ["许可", "CC BY-NC-SA 4.0，仅用于本次非商业科研/竞赛评测"],
        ],
        [2400, 6960],
        font_size=8.7,
    )

    add_heading(doc, "3.2 运行环境与冻结契约", 2)
    add_table_caption(doc, "表 3  本地离线扫描环境")
    add_data_table(
        doc,
        ["项目", "设置"],
        [
            ["操作系统", "Windows 11，本地 CPU，顺序扫描"],
            ["Python", "3.13.14"],
            ["Cisco Skill Scanner", "2.0.13.dev3+g4dee90371"],
            ["策略", "aegis-chain-local-default 1.0.0"],
            ["单条超时", "150 秒"],
            ["允许分析器", "static_analyzer / bytecode / pipeline"],
            ["LLM、云上传、行为执行", "全部关闭"],
            ["总墙钟时间", "约 39 分钟，包含中断定位与断点恢复"],
        ],
        [2800, 6560],
    )

    add_heading(doc, "3.3 三套评价口径", 2)
    add_body(doc, "严格三分类：normal、suspicious、malicious 分别对应 ALLOW、REVIEW、BLOCK；UNKNOWN 为 abstain，并在主指标中按错误计。")
    add_body(doc, "风险筛查二分类：suspicious 与 malicious 合并为 non-normal；REVIEW 或 BLOCK 均算识别到风险，衡量接入门禁能否把风险样本送入人工复核或阻断。")
    add_body(doc, "运营失败闭锁：REVIEW、BLOCK、UNKNOWN 均不自动放行。该口径说明系统是否安全处置失败，不能与检测模型 F1 混为一谈。")

    add_heading(doc, "4 扫描结果与指标解读", 1)
    add_heading(doc, "4.1 运行完整性", 2)
    add_callout(
        doc,
        "运行终态",
        "556/556 条样本均产生终态：546 条 completed，7 条 Cisco 运行时错误，3 条被 Windows Defender 阻断读取；后 10 条全部记为 UNKNOWN/abstain。",
        fill=PALE_GREEN,
        accent=GREEN,
    )
    add_table_caption(doc, "表 4  核心指标")
    add_data_table(
        doc,
        ["指标", "结果", "正确解释"],
        [
            ["Coverage", "98.20%", "546/556 获得非 abstain 预测"],
            ["Failure rate", "1.80%", "10/556 进入失败闭锁"],
            ["Strict macro F1", "0.4977", "三类 F1 等权平均，abstain 计错"],
            ["三分类准确率", "60.61%", "仅作补充，不作为单一宣传指标"],
            ["Malicious recall", "72.28%", "285 条恶意中 206 条被 BLOCK"],
            ["Non-normal recall", "76.67%", "390 条风险样本中 299 条进入 REVIEW/BLOCK"],
            ["Normal FPR", "24.70%", "166 条正常中 41 条被 REVIEW/BLOCK"],
            ["中位 / P95 耗时", "4,028 / 4,309 ms", "单机顺序静态扫描，耗时稳定"],
        ],
        [2350, 1700, 5310],
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    core_picture = p.add_run().add_picture(str(core_chart), width=Inches(6.15))
    core_picture._inline.docPr.set("title", "核心指标条形图")
    core_picture._inline.docPr.set("descr", "覆盖率98.20%，风险筛查精确率87.94%，风险筛查召回率76.67%，恶意严格召回率72.28%，正常样本误报率24.70%。")
    add_table_caption(doc, "图 1  核心运行与风险筛查指标（严格 macro F1 另按三分类定义计算）")

    add_heading(doc, "4.2 严格三分类与门禁二分类的差异", 2)
    add_body(
        doc,
        "Strict macro F1 为 0.4977，并不意味着系统只能发现约一半风险。该指标要求三类精确区分：例如 malicious 被判为 REVIEW 时，在安全运营中已经被拦下，但在严格三分类中仍然算错。suspicious 类只有 9/105 被精确判为 suspicious，是 macro F1 偏低的主要原因之一。",
    )
    add_body(
        doc,
        "Loose F1 为 81.92%，回答的是另一个问题：系统能否把 suspicious 或 malicious 样本送入 REVIEW/BLOCK。它更接近供应链准入门的实际工作方式，但不能替代严格三分类指标。正式汇报必须同时给出两者，避免只挑选更好看的数字。",
    )

    add_table_caption(doc, "表 5  统一策略层二分类结果")
    add_data_table(
        doc,
        ["TP", "FP", "FN", "TN", "Precision", "Recall", "Loose F1", "FPR"],
        [["299", "41", "91", "122", "87.94%", "76.67%", "81.92%", "24.70%"]],
        [900, 900, 900, 900, 1440, 1440, 1440, 1440],
        alignments=[WD_ALIGN_PARAGRAPH.CENTER] * 8,
        font_size=9,
    )
    add_body(
        doc,
        "注：风险侧 7 个 abstain 按未检出风险计入 FN；正常侧 3 个 abstain 不计为 FP，也不计为 TN，但仍保留在 166 条 normal 的 FPR 分母中。因此 TP+FP+FN+TN 与独立列出的 abstain 不是互斥分组。",
    )

    add_heading(doc, "4.3 三分类混淆矩阵", 2)
    add_table_caption(doc, "表 6  三分类混淆矩阵（行是真值，列是系统预测）")
    add_data_table(
        doc,
        ["Ground truth", "Normal", "Suspicious", "Malicious", "Abstain", "合计"],
        [
            ["Normal", "122", "22", "19", "3", "166"],
            ["Suspicious", "38", "9", "57", "1", "105"],
            ["Malicious", "46", "27", "206", "6", "285"],
        ],
        [2100, 1452, 1452, 1452, 1452, 1452],
        alignments=[WD_ALIGN_PARAGRAPH.LEFT] + [WD_ALIGN_PARAGRAPH.CENTER] * 5,
    )
    add_callout(
        doc,
        "矩阵解读",
        "主要问题不是完全检不出恶意，而是精细分级能力不足：27 条恶意被降为 REVIEW，安全门禁仍拦截；46 条恶意被 ALLOW，属于真正需要优先处理的高风险漏报。",
        fill=PALE_GOLD,
        accent=GOLD,
    )

    add_heading(doc, "4.4 T01–T09 风险类型召回", 2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    risk_picture = p.add_run().add_picture(str(risk_chart), width=Inches(6.15))
    risk_picture._inline.docPr.set("title", "T01至T09风险召回率柱状图")
    risk_picture._inline.docPr.set("descr", "T01至T09召回率依次为82.31%、75.00%、76.92%、79.24%、66.67%、40.00%、91.67%、96.00%、62.50%，其中T06、T09、T05为补强重点。")
    add_table_caption(doc, "图 2  T01–T09 风险筛查召回率；T06、T09、T05 为当前补强重点")

    add_data_table(
        doc,
        ["风险类型", "Support", "Detected", "Recall"],
        [
            ["T01", "147", "121", "82.31%"], ["T02", "16", "12", "75.00%"],
            ["T03", "78", "60", "76.92%"], ["T04", "289", "229", "79.24%"],
            ["T05", "111", "74", "66.67%"], ["T06", "10", "4", "40.00%"],
            ["T07", "12", "11", "91.67%"], ["T08", "25", "24", "96.00%"],
            ["T09", "112", "70", "62.50%"],
        ],
        [3000, 1900, 1900, 2560],
        alignments=[WD_ALIGN_PARAGRAPH.CENTER] * 4,
        font_size=9.2,
    )
    add_body(doc, "说明：一条样本可包含多个风险标签，因此各类型 support 之和大于 390。T06、T07 support 较小，百分比波动较大；T06 的 4/10 仍足以作为高优先级人工复核切片，但不宜把单个百分比外推到真实市场。")

    add_heading(doc, "5 误报、漏报与失败分析", 1)
    add_heading(doc, "5.1 风险筛查漏报", 2)
    add_body(doc, "共 91 条 non-normal 未进入 REVIEW/BLOCK，其中 39 条 suspicious、52 条 malicious。严格恶意漏判共 79 条：27 条进入 REVIEW，46 条被 ALLOW，6 条为 abstain。")
    add_table_caption(doc, "表 7  高优先级漏报切片")
    add_data_table(
        doc,
        ["切片", "现象", "风险判断", "下一步"],
        [
            ["wild_real_world", "风险筛查仅 6/27，召回 22.22%", "最明显的分布外短板", "逐案提取跨文件、隐蔽指令和真实攻击模式"],
            ["crypto_wallet", "风险筛查 16/27，召回 59.26%", "涉及密钥、签名和外联，业务与攻击语义接近", "增加敏感数据来源—目标域—网络行为关联"],
            ["api_integration", "15 条漏报", "合法 API 调用掩盖潜在外传路径", "结合声明能力、目标域和数据类型做上下文判定"],
            ["T06/T09/T05", "召回 40.00% / 62.50% / 66.67%", "规则覆盖不足或严重度门槛不合适", "每类选择 5–10 条代表样本提炼规则"],
        ],
        [1750, 2460, 2600, 2550],
        font_size=8.6,
    )

    add_heading(doc, "5.2 正常样本误报", 2)
    add_body(
        doc,
        "166 条 normal 中有 41 条被判为 non-normal：22 条 REVIEW、19 条 BLOCK。较常见的触发包括 DATA_EXFIL_JS_FS_ACCESS、TOOL_ABUSE_UNDECLARED_NETWORK、DATA_EXFIL_NETWORK_REQUESTS 和 FILE_MAGIC_MISMATCH。网络、文件和编码行为既可能是正常工具能力，也可能是恶意外传手段，因此不能简单下调所有严重度。",
    )
    add_callout(
        doc,
        "改进原则",
        "把“声明了网络能力”和“未声明却调用网络”分开处理，并关联目标域、数据来源、访问范围和输出去向；在上下文证据不足时优先 REVIEW，而不是全局放宽为 ALLOW。",
        fill=PALE_BLUE,
        accent=BLUE,
    )

    add_heading(doc, "5.3 运行失败与安全事件", 2)
    add_body(doc, "7 条 Cisco 运行时错误分别涉及 3 条 malicious、1 条 suspicious 和 3 条 normal；当前未保留可能包含样本内容的原始错误文本，主结果不重试、不启用 lenient frontmatter，也不覆盖为安全。")
    add_body(doc, "case_05527、case_05559、case_05568 三条 malicious 被 Windows Defender 阻断读取或隔离。项目没有关闭 Defender、没有设置排除目录、没有重新释放样本，均按 UNKNOWN/abstain 记入结果。")
    add_body(doc, "首次运行在第 529 条后停止。恢复前重新校验运行 ID、样本顺序、数据 revision、清单/扫描器/策略/指标契约哈希，以及前 529 条逐样本 before/after hash；全部一致后才从第 530 条继续。")

    add_heading(doc, "6 与 90 条 Pilot 基线的对照", 1)
    add_body(doc, "90 条 pilot 为三类各 30 条的平衡工程样本；官方 556 条子集为 166/105/285，并包含 wild real-world 样本。两者分布不同，差值只能用于观察工程稳定性与风险方向，不能解释为同分布性能变化。")
    add_table_caption(doc, "表 8  Pilot90 与官方 556 条子集对照")
    add_data_table(
        doc,
        ["指标", "Pilot90", "官方 556", "变化"],
        [
            ["Coverage", "98.89%", "98.20%", "-0.69 pp"],
            ["Failure rate", "1.11%", "1.80%", "+0.69 pp"],
            ["Strict macro F1", "0.5114", "0.4977", "-0.0137"],
            ["Malicious recall", "80.00%", "72.28%", "-7.72 pp"],
            ["Non-normal recall", "78.33%", "76.67%", "-1.67 pp"],
            ["Normal FPR", "33.33%", "24.70%", "-8.63 pp"],
            ["Loose precision", "82.46%", "87.94%", "+5.49 pp"],
            ["Loose F1", "80.34%", "81.92%", "+1.58 pp"],
            ["中位耗时", "3,935 ms", "4,028 ms", "+93 ms"],
            ["P95 耗时", "4,226 ms", "4,309 ms", "+83 ms"],
        ],
        [3300, 2020, 2020, 2020],
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER],
        font_size=9.2,
    )
    add_body(doc, "稳定结论是：覆盖率仍约 98%，单条耗时仍约 4 秒，strict macro F1 仍约 0.50，策略层 loose F1 仍约 0.80。需要修正的是，90 条 pilot 对恶意严格召回偏乐观，扩大到官方子集后降至 72.28%，并暴露 wild real-world 的明显短板。")

    add_heading(doc, "7 与官方 Cisco 榜单的关系", 1)
    add_table_caption(doc, "表 9  外部参照（不可直接等价比较）")
    add_data_table(
        doc,
        ["口径", "Precision", "Recall", "Loose F1", "FPR"],
        [
            ["官方 Cisco 行", "90.07%", "95.38%", "92.65%", "24.70%"],
            ["本地 Aegis 策略层", "87.94%", "76.67%", "81.92%", "24.70%"],
        ],
        [2960, 1600, 1600, 1600, 1600],
        alignments=[WD_ALIGN_PARAGRAPH.LEFT] + [WD_ALIGN_PARAGRAPH.CENTER] * 4,
    )
    add_body(doc, "不能称为“复现官方 Cisco 成绩”，原因如下：")
    for text in (
        "官方 Cisco 行直接使用扫描器 actual_safe 二分类；本轮经过 Aegis Chain 严重度与准入策略映射。",
        "官方工具比较行标注 DeepSeek v4 Flash；本轮完全关闭 LLM 与云分析。",
        "扫描器版本、运行环境和失败处理方式未证明一致；本轮 10 个 abstain 在严格召回中按错误计。",
        "FPR 四舍五入后相同只说明正常样本误报比例接近，不能证明整体流程和得分等价。",
    ):
        add_bullet(doc, text)
    add_callout(
        doc,
        "推荐表述",
        "在 SkillTrustBench 官方 556 条固定清单上，完成了 Cisco Skill Scanner 的本地离线接入与 Aegis Chain 统一策略层评测；结果可复算、可追溯，但评价口径与官方 Cisco 榜单不同。",
        fill=PALE_GREEN,
        accent=GREEN,
    )

    add_heading(doc, "8 当前成果与项目创新点", 1)
    add_heading(doc, "8.1 已完成成果", 2)
    for text in (
        "完成 Cisco Skill Scanner 本地静态底座接入，并验证 Cisco MCP Scanner 的可运行性与适用边界。",
        "建立统一 Finding 与 ALLOW/REVIEW/BLOCK/UNKNOWN 四态策略契约，避免厂商原始分数直接变成系统最终结论。",
        "建立数据安全导入、版本/哈希追踪、批量运行、超时、失败闭锁、断点续扫和脱敏证据输出。",
        "完成 90 条 pilot 与官方 556 条固定子集的两阶段评测；指标、混淆矩阵和错误切片可由逐样本结果复算。",
        "后端自动测试达到 73 passed，为继续迭代提供回归保护。",
    ):
        add_bullet(doc, text)

    add_heading(doc, "8.2 可用于竞赛表达的创新", 2)
    add_body(doc, "本项目的创新不在于“调用了 Cisco 工具”，而在于把第三方检测能力转化为政企可用的安全工程链路：")
    for text in (
        "多工具适配边界：第三方原始输出与统一策略解耦，后续可接入第二扫描器并做互补证据融合。",
        "四态安全决策：显式保留 UNKNOWN，使工具失败不等于安全，符合政企场景的失败闭锁要求。",
        "可复现证据链：固定数据、工具、策略和指标版本，支持断点恢复与逐样本哈希验证。",
        "静态—动态分层：静态阶段负责大规模低成本筛查，动态阶段只对重点样本提供行为证据，兼顾安全、成本与解释性。",
        "面向治理的输出：既报告检测能力，也报告覆盖率、失败率、误报、漏报和不可比较边界，避免只给单一高分。",
    ):
        add_bullet(doc, text)

    add_heading(doc, "9 8 月底前开发计划", 1)
    add_table_caption(doc, "表 10  后续两周实施计划")
    add_data_table(
        doc,
        ["时间", "优先级", "主要任务", "可验收产物"],
        [
            ["8 月 15–18 日", "P0", "人工复核 91 条 risk-screening miss，优先 wild_real_world、crypto_wallet 和 T06/T09/T05", "错误分类表、代表样本证据链、候选规则清单"],
            ["8 月 19–22 日", "P0/P1", "实现 2–3 组上下文规则，覆盖声明能力、数据来源、目标域与行为组合", "规则配置、单元测试、冻结前开发结果"],
            ["8 月 23–26 日", "P1", "用无害 fixture 建立进程创建、网络连接、文件写入、敏感环境变量访问四类动态事件", "最小动态监控原型、事件 JSON、演示脚本"],
            ["8 月 27–29 日", "P1", "统一静态/动态结果接口，完成平台接入适配；若统一平台仍未提供，保留独立 Web 演示", "API 契约、可演示页面、样例报告"],
            ["8 月 30–31 日", "P2", "冻结策略并复测，整理最终材料与演示容错", "对照指标、操作手册、最终汇报材料"],
        ],
        [1500, 1100, 4260, 2500],
        font_size=8.4,
    )

    add_heading(doc, "9.1 验收标准", 2)
    for text in (
        "静态主链路可对单个 Skill 和批量样本稳定输出四态结论、Finding 和可追溯元数据。",
        "新增规则必须有对应正常/异常单元测试，并在冻结评测上报告召回、误报与回归变化。",
        "动态阶段只使用自建无害模拟样本验证监控链路，不直接运行 SkillTrustBench 真实恶意样本。",
        "演示时能展示一次正常通过、一次人工复核、一次阻断和一次扫描失败闭锁。",
        "最终汇报明确说明数据来源、许可、指标口径、限制与下一步，不夸大为官方复现或最终产品能力。",
    ):
        add_bullet(doc, text)

    add_heading(doc, "10 阶段结论", 1)
    add_body(
        doc,
        "本阶段已经完成“能不能运行”到“能不能被可信评估”的跨越。Cisco Skill Scanner 作为静态底座，在官方 556 条子集上保持约 98% 覆盖率和约 4 秒单条耗时；统一门禁 loose F1 为 81.92%，说明其具备现实筛查价值。与此同时，严格三分类 macro F1 为 0.4977、恶意严格召回率为 72.28%，以及 wild real-world 22.22% 的风险筛查召回，清楚表明该底座还不能承担最终安全裁决。",
    )
    add_body(
        doc,
        "因此，项目后续不应追求简单堆叠更多工具或只提高单一分数，而应围绕已定位的错误切片补充上下文规则，并用最小动态证据验证静态难以判断的行为。最终作品的合理定位是：一个面向政企智能体工具接入的“供应链安全准入与审计模块”，能够统一接入第三方扫描器、执行失败闭锁、输出可解释证据，并为动态审计和统一平台治理提供标准接口。",
    )

    doc.add_page_break()
    add_heading(doc, "附录 A  指标定义与技术解释", 1)
    definitions = [
        ("Coverage（覆盖率）", "完成非 abstain 预测的样本比例。本轮为 546/556=98.20%。它衡量工程可用性，不等于检测正确率。"),
        ("Failure rate（失败率）", "进入 UNKNOWN/abstain 的样本比例。本轮为 10/556=1.80%。失败闭锁保证异常不会被静默当作安全。"),
        ("Strict macro F1", "分别计算 normal、suspicious、malicious 三类 F1 后等权平均；类别样本量不同也给予相同权重，abstain 按错误计。本轮为 0.4977。"),
        ("Malicious recall", "被正确判为 BLOCK 的 malicious 占全部 malicious 的比例：206/285=72.28%。REVIEW 在该严格指标中仍算错误。"),
        ("Malicious FNR", "恶意样本未被严格判为 BLOCK 的比例：79/285=27.72%，与 malicious recall 互补。"),
        ("Non-normal recall", "suspicious/malicious 中进入 REVIEW 或 BLOCK 的比例：299/390=76.67%。它更贴近准入门筛查能力。"),
        ("Precision", "所有 REVIEW/BLOCK 中真实 non-normal 的比例：299/(299+41)=87.94%。"),
        ("Normal FPR", "normal 被错误送入 REVIEW/BLOCK 的比例：41/166=24.70%。本项目分母包含 3 个 normal abstain，以完整 normal 集为基准。"),
        ("Loose F1", "风险筛查 precision 与 recall 的调和平均：2PR/(P+R)=81.92%。它不能替代三分类 macro F1。"),
        ("P95 耗时", "95% 已完成样本的单条耗时不超过该值。本轮为 4,309 ms，用于判断演示和批量处理的时延稳定性。"),
    ]
    for term, definition in definitions:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(7)
        r = p.add_run(term + "：")
        set_run_font(r, east_asia_name=EAST_ASIA_HEAD, bold=True, color=NAVY)
        r = p.add_run(definition)
        set_run_font(r)

    add_heading(doc, "附录 B  5 分钟正式汇报稿", 1)
    speech_paragraphs = [
        "各位老师好，我负责 XA-202620 赛题中的供应链安全模块。本模块关注的问题是：大模型智能体接入 Skill、MCP Server 和插件后，如何在使用前发现其中可能存在的恶意指令、隐蔽网络访问、敏感数据读取和依赖风险，并形成可解释、可追溯的准入结论。",
        "现阶段我选择 Cisco Skill Scanner 作为 Skill 静态审查底座，同时验证了 Cisco MCP Scanner 的可运行性。两者用途不同：Skill Scanner 用于审查 Skill 文件与代码，MCP Scanner 面向 MCP server 或配置。因此本轮先完成 Skill 侧主链路，没有为了展示效果把 MCP Scanner 强行套用到不匹配的数据集。",
        "为了避免只用少量自造样本，我选用了腾讯朱雀实验室与香港中文大学（深圳）联合发布的 SkillTrustBench，并严格固定官方 10% 子集 556 条，其中 normal 166 条、suspicious 105 条、malicious 285 条。整个实验在本地离线完成，关闭 LLM、云上传和行为执行，不安装、不导入、不执行任何样本代码。每条样本在扫描前后都校验哈希，保证评测过程没有修改样本。",
        "结果方面，556 条样本全部得到终态，其中 546 条完成 Cisco 扫描，7 条出现运行时错误，3 条被 Windows Defender 阻断读取。这 10 条没有被当作安全，而是统一记为 UNKNOWN，也就是失败闭锁。最终覆盖率为 98.20%，单条中位耗时约 4 秒，说明静态链路已经具备批量运行和演示基础。",
        "能力指标需要分两套理解。第一套是严格三分类：必须准确区分 normal、suspicious 和 malicious，macro F1 为 0.4977，恶意严格召回率为 72.28%。第二套是实际准入门的风险筛查：只要风险样本进入 REVIEW 或 BLOCK，就认为被发现。该口径 precision 为 87.94%，recall 为 76.67%，loose F1 为 81.92%。这两个数字不矛盾：例如恶意样本被判为 REVIEW，安全上已经被拦住，但三分类标签仍然判错。",
        "错误分析显示，目前有 91 条风险样本未进入 REVIEW 或 BLOCK，最明显的短板是 wild real-world，召回只有 6/27；crypto_wallet 为 16/27。风险类型上 T06、T09、T05 召回较低。正常样本中有 41 条被误报，常见原因是网络请求和文件访问缺少上下文，同一种 API 既可能是合法业务能力，也可能用于数据外传。",
        "因此我的下一步不是简单堆叠工具，而是先处理这 91 条漏报，提炼 2 到 3 组可解释的上下文规则，把声明能力、实际行为、目标域和数据类型关联起来。随后用自建无害样本建立进程、网络、文件写入和敏感环境变量访问四类动态事件，再把静态和动态证据统一输出。",
        "本阶段的结论是：Cisco 工具已经成为可复现的静态底座，但不能直接作为最终裁决。本项目的工作价值在于统一证据模型、四态决策、失败闭锁、版本与哈希追踪，以及静态到动态的分层审计。到 8 月底，我将完成规则补强、最小动态闭环、平台接口和最终演示材料。谢谢各位老师。",
    ]
    for paragraph in speech_paragraphs:
        add_body(doc, paragraph)

    add_heading(doc, "附录 C  常见答辩问题与建议回答", 1)
    qa = [
        ("为什么 strict macro F1 只有 0.4977，还能说系统有价值？", "因为该指标要求三类精确分级，尤其 suspicious 只有 9/105 被精确命中；但供应链门禁更关心风险是否进入 REVIEW/BLOCK，二分类 loose F1 为 81.92%。我同时报告两套口径，不回避分级短板，也不把筛查能力夸大成最终判定能力。"),
        ("为什么只测 556 条，不跑完整 5,520 条？", "556 条不是随意抽样，而是官方结果仓库固定的 10% 公开评测清单，便于与公开工具结果在相同 ID 上参照；在单机与时间约束下，它也能保证结果可复现。完整集可在系统冻结后作为扩展评测。"),
        ("为什么不关闭 Windows Defender，把 3 条样本继续扫完？", "关闭宿主机防护会扩大真实恶意样本的暴露风险，也破坏安全工程原则。我把这 3 条记为 UNKNOWN 并禁止自动放行，保留固定归档哈希作为身份；后续若必须处理，应使用专门隔离沙箱，而不是降低本机防护。"),
        ("和官方 Cisco 榜单为什么差距较大？", "官方行使用 actual_safe 二分类并标注 DeepSeek v4 Flash，本轮是本地离线、无 LLM，并经过 Aegis 策略映射，工具版本和失败处理也未证明一致。因此只能说在相同 556 条清单上完成本地策略层评测，不能称为复现官方分数。"),
        ("项目创新是不是只是集成开源工具？", "第三方扫描器只是底座。项目创新在于把多工具输出适配为统一 Finding，建立 ALLOW/REVIEW/BLOCK/UNKNOWN 四态准入、失败闭锁、版本/哈希证据链、批量可复现评测，以及后续静态—动态分层融合。"),
        ("下一阶段最关键的可量化目标是什么？", "一是围绕 91 条 risk-screening miss 做分组分析并实现 2–3 组上下文规则；二是在固定评测上同时报告召回提升和 FPR 变化；三是用无害 fixture 完成四类动态行为事件的采集与统一输出。"),
    ]
    for question, answer in qa:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(3)
        set_paragraph_keep(p)
        r = p.add_run("问：" + question)
        set_run_font(r, east_asia_name=EAST_ASIA_HEAD, color=NAVY, bold=True)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run("答：" + answer)
        set_run_font(r)

    add_heading(doc, "参考资料与项目证据", 1)
    reference_num_id = create_decimal_numbering(doc)
    p = add_number(doc, "", reference_num_id, keep_next=True)
    r = p.add_run("SkillTrustBench 数据集：")
    set_run_font(r)
    add_hyperlink(p, "访问官方数据集页面", "https://huggingface.co/datasets/cuhk-zhuque/SkillTrustBench")
    p = add_number(doc, "", reference_num_id, keep_next=True)
    r = p.add_run("SkillTrustBench 官方结果与固定 10% 清单：")
    set_run_font(r)
    add_hyperlink(p, "访问官方结果页面", "https://huggingface.co/datasets/cuhk-zhuque/SkillTrustBench-results")
    add_number(doc, "项目内部证据：M2_SKILLTRUSTBENCH_OFFICIAL_10PCT_REPORT.md、运行清单、逐 case 结果、metrics.json、confusion_matrix.json 与错误清单。", reference_num_id)

    add_callout(
        doc,
        "材料使用建议",
        "老师快速阅读时优先看摘要、表 4、表 6、图 2、表 10 和阶段结论；现场 5 分钟汇报可直接使用附录 B；追问指标或技术边界时查阅附录 A、附录 C。",
        fill=PALE_BLUE,
        accent=BLUE,
    )

    # Re-apply geometry after all content is present and ensure no row uses fixed height.
    for table in doc.tables:
        for row in table.rows:
            tr_pr = row._tr.get_or_add_trPr()
            tr_height = tr_pr.find(qn("w:trHeight"))
            if tr_height is not None:
                tr_pr.remove(tr_height)

    doc.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()
    build_report(args.output.resolve())
    print(args.output.resolve())


if __name__ == "__main__":
    main()
