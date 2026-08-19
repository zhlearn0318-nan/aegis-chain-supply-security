from __future__ import annotations

from datetime import date
from pathlib import Path
from typing import Iterable, Sequence

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs"
ASSET_DIR = OUT_DIR / "_report_assets"
OUTPUT = OUT_DIR / "Agent供应链安全静态扫描与来源证明开发参考报告.docx"

FONT_REGULAR = Path(r"C:\Windows\Fonts\msyh.ttc")
FONT_BOLD = Path(r"C:\Windows\Fonts\msyhbd.ttc")

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "0B2545"
MUTED = "5F6B76"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"
GREEN = "1F6B52"
AMBER = "8A6500"
RED = "9B1C1C"
BLACK = "111111"

PAGE_WIDTH_DXA = 12240
PAGE_HEIGHT_DXA = 15840
MARGIN_DXA = 1440
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_dxa: Sequence[int], indent_dxa: int = TABLE_INDENT_DXA) -> None:
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ValueError(f"Column widths must sum to {CONTENT_WIDTH_DXA}: {widths_dxa}")
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[idx]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color="CDD3DA", size="6") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_run_font(
    run,
    *,
    latin="Calibri",
    east_asia="Microsoft YaHei",
    size: float | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    color: str | None = None,
) -> None:
    run.font.name = latin
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), latin)
    r_fonts.set(qn("w:hAnsi"), latin)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_style_font(style, size: float, color: str = BLACK, bold: bool = False) -> None:
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_fonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    set_style_font(normal, 11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    h1 = doc.styles["Heading 1"]
    set_style_font(h1, 16, BLUE, True)
    h1.paragraph_format.space_before = Pt(16)
    h1.paragraph_format.space_after = Pt(8)
    h1.paragraph_format.keep_with_next = True

    h2 = doc.styles["Heading 2"]
    set_style_font(h2, 13, BLUE, True)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True

    h3 = doc.styles["Heading 3"]
    set_style_font(h3, 12, DARK_BLUE, True)
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.keep_with_next = True

    caption = doc.styles["Caption"]
    set_style_font(caption, 9.5, MUTED, False)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.keep_with_next = True


def add_numbering_definitions(doc: Document) -> tuple[int, int]:
    numbering = doc.part.numbering_part.element
    existing_abs = [
        int(n.get(qn("w:abstractNumId")))
        for n in numbering.findall(qn("w:abstractNum"))
        if n.get(qn("w:abstractNumId")) is not None
    ]
    existing_num = [
        int(n.get(qn("w:numId")))
        for n in numbering.findall(qn("w:num"))
        if n.get(qn("w:numId")) is not None
    ]
    next_abs = max(existing_abs, default=0) + 1
    next_num = max(existing_num, default=0) + 1

    def create(kind: str, abstract_id: int, num_id: int) -> None:
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
        lvl.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
        lvl.append(lvl_text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        lvl.append(suff)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "720")
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "720")
        ind.set(qn("w:hanging"), "360")
        p_pr.append(ind)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "160")
        spacing.set(qn("w:line"), "280")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.append(spacing)
        lvl.append(p_pr)
        r_pr = OxmlElement("w:rPr")
        r_fonts = OxmlElement("w:rFonts")
        r_fonts.set(qn("w:ascii"), "Calibri")
        r_fonts.set(qn("w:hAnsi"), "Calibri")
        r_fonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        r_pr.append(r_fonts)
        lvl.append(r_pr)
        abstract.append(lvl)
        numbering.append(abstract)

        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abs_id = OxmlElement("w:abstractNumId")
        abs_id.set(qn("w:val"), str(abstract_id))
        num.append(abs_id)
        numbering.append(num)

    create("bullet", next_abs, next_num)
    create("decimal", next_abs + 1, next_num + 1)
    return next_num, next_num + 1


def add_numbered_paragraph(doc: Document, text: str, num_id: int, bold_prefix: str | None = None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)
    p_pr.append(num_pr)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2)
    else:
        run = p.add_run(text)
        set_run_font(run)
    return p


def add_body(doc: Document, text: str, *, bold_prefix: str | None = None, keep_with_next=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    p.paragraph_format.keep_with_next = keep_with_next
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_callout(doc: Document, label: str, text: str, fill=CALLOUT, accent=BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    set_table_borders(table, color=accent, size="10")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.10
    r1 = p.add_run(f"{label}  ")
    set_run_font(r1, bold=True, color=accent)
    r2 = p.add_run(text)
    set_run_font(r2, color=BLACK)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_table(
    doc: Document,
    headers: Sequence[str],
    rows: Sequence[Sequence[str]],
    widths: Sequence[int],
    *,
    font_size: float = 9.2,
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_GRAY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_run_font(r, size=font_size, bold=True, color=NAVY)
    for row_values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            cell = cells[idx]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            r = p.add_run(str(value))
            set_run_font(r, size=font_size)
    after = doc.add_paragraph()
    after.paragraph_format.space_before = Pt(2)
    after.paragraph_format.space_after = Pt(2)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = paragraph.add_run("第 ")
    set_run_font(r1, size=9, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    fld.append(r)
    paragraph._p.append(fld)
    r2 = paragraph.add_run(" 页")
    set_run_font(r2, size=9, color=MUTED)


def configure_section(section) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Agent 供应链安全｜阶段设计参考")
    set_run_font(r, size=9, color=MUTED, bold=True)

    footer = section.footer
    fp = footer.paragraphs[0]
    add_page_number(fp)


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.FreeTypeFont, max_width: int) -> list[str]:
    lines: list[str] = []
    for part in text.split("\n"):
        current = ""
        for ch in part:
            trial = current + ch
            if draw.textlength(trial, font=font) <= max_width:
                current = trial
            else:
                if current:
                    lines.append(current)
                current = ch
        lines.append(current)
    return lines


def draw_centered_text(draw, box, text, font, fill, max_width=None, spacing=6):
    x1, y1, x2, y2 = box
    max_width = max_width or (x2 - x1 - 30)
    lines = wrap_text(draw, text, font, max_width)
    line_height = font.size + spacing
    total = len(lines) * line_height - spacing
    y = y1 + (y2 - y1 - total) / 2
    for line in lines:
        width = draw.textlength(line, font=font)
        draw.text((x1 + (x2 - x1 - width) / 2, y), line, font=font, fill=fill)
        y += line_height


def arrow(draw, start, end, color=NAVY, width=5):
    color = f"#{color}" if isinstance(color, str) and len(color) == 6 and not color.startswith("#") else color
    draw.line([start, end], fill=color, width=width)
    x2, y2 = end
    x1, y1 = start
    dx, dy = x2 - x1, y2 - y1
    length = max((dx * dx + dy * dy) ** 0.5, 1)
    ux, uy = dx / length, dy / length
    px, py = -uy, ux
    size = 15
    p1 = (x2, y2)
    p2 = (x2 - ux * size + px * size * 0.55, y2 - uy * size + py * size * 0.55)
    p3 = (x2 - ux * size - px * size * 0.55, y2 - uy * size - py * size * 0.55)
    draw.polygon([p1, p2, p3], fill=color)


def build_architecture_diagram(path: Path) -> None:
    pil = lambda color: f"#{color}" if isinstance(color, str) and len(color) == 6 and not color.startswith("#") else color
    img = Image.new("RGB", (1800, 960), pil(WHITE))
    draw = ImageDraw.Draw(img)
    title = ImageFont.truetype(str(FONT_BOLD), 44)
    body = ImageFont.truetype(str(FONT_REGULAR), 31)
    small = ImageFont.truetype(str(FONT_REGULAR), 25)
    draw.text((70, 45), "建议系统架构：Cisco 作为检测执行器，自研层负责可信决策", font=title, fill=pil(NAVY))

    boxes = [
        ((80, 180, 340, 380), "输入制品\nSkill / MCP\n依赖包", LIGHT_BLUE),
        ((430, 150, 750, 410), "来源证明\n仓库、提交、签名\n文件哈希、SBOM\n只读快照", CALLOUT),
        ((850, 105, 1210, 285), "Skill Scanner\nStatic / YARA\nBytecode / Pipeline", LIGHT_BLUE),
        ((850, 330, 1210, 510), "MCP Scanner\nTool / Prompt / Resource\nYARA / pip-audit", LIGHT_BLUE),
        ((1320, 180, 1690, 430), "统一 Finding IR\n证据、严重度、状态\n扫描器版本、策略指纹\nSAFE / UNSAFE / UNKNOWN", CALLOUT),
    ]
    for box, label, fill in boxes:
        draw.rounded_rectangle(box, radius=24, fill=pil(fill), outline=pil(BLUE), width=4)
        draw_centered_text(draw, box, label, body, pil(NAVY))

    arrow(draw, (340, 280), (430, 280))
    arrow(draw, (750, 235), (850, 195))
    arrow(draw, (750, 325), (850, 420))
    arrow(draw, (1210, 195), (1320, 265))
    arrow(draw, (1210, 420), (1320, 345))

    future_box = (220, 640, 1580, 860)
    draw.rounded_rectangle(future_box, radius=24, fill="#FFF8E8", outline="#C08A00", width=4)
    future_title = ImageFont.truetype(str(FONT_BOLD), 32)
    draw.text((270, 675), "后续扩展层", font=future_title, fill=pil(AMBER))
    items = [
        "语义证据复核",
        "跨 Skill / MCP 证据图",
        "组合攻击链分析",
        "隔离沙箱与动态验证",
    ]
    item_w = 280
    gap = 35
    start_x = 260
    for idx, item in enumerate(items):
        box = (start_x + idx * (item_w + gap), 745, start_x + idx * (item_w + gap) + item_w, 825)
        draw.rounded_rectangle(box, radius=16, fill=pil(WHITE), outline="#D9B85C", width=3)
        draw_centered_text(draw, box, item, small, pil(NAVY))
    path.parent.mkdir(parents=True, exist_ok=True)
    img.save(path, quality=95)


def add_title_block(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("技术设计参考报告")
    set_run_font(r, size=11, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("Agent 供应链安全平台")
    set_run_font(r, size=25, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run("静态规则扫描、语义检测与来源证明的集成方案")
    set_run_font(r, size=15, color=DARK_BLUE)

    metadata = [
        ("报告用途", "下一阶段开发设计、结论整理与验收参考"),
        ("技术底座", "Cisco Skill Scanner + Cisco MCP Scanner"),
        ("当前范围", "静态检测、语义复核、来源证明、统一风险门禁"),
        ("后续范围", "跨组件组合攻击、隔离沙箱、动态数据流验证"),
        ("版本日期", str(date.today())),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(f"{label}：")
        set_run_font(r1, size=10.5, bold=True, color=BLACK)
        r2 = p.add_run(value)
        set_run_font(r2, size=10.5, color=BLACK)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(12)
    p_pr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "16")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), BLUE)
    borders.append(bottom)
    p_pr.append(borders)


def add_section_map(doc: Document, bullet_num: int) -> None:
    doc.add_heading("报告结构", level=1)
    items = [
        "第1—3章：项目目标、范围与两个 Cisco 底座的选型结论。",
        "第4—8章：总体架构、静态规则、语义检测、来源证明和统一结果模型。",
        "第9—12章：风险门禁、工程集成、评测体系和阶段开发计划。",
        "附录：字段样例、接口约定、开发验收清单与参考来源。",
    ]
    for item in items:
        add_numbered_paragraph(doc, item, bullet_num)


def build_document() -> Path:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    architecture = ASSET_DIR / "target_architecture.png"
    build_architecture_diagram(architecture)

    doc = Document()
    configure_styles(doc)
    for section in doc.sections:
        configure_section(section)
    bullet_num, decimal_num = add_numbering_definitions(doc)

    props = doc.core_properties
    props.title = "Agent供应链安全静态扫描与来源证明开发参考报告"
    props.subject = "Cisco Skill Scanner 与 MCP Scanner 集成设计"
    props.author = "项目技术研究与开发参考"
    props.keywords = "Agent, Skill, MCP, 供应链安全, 静态扫描, 来源证明, 语义检测"

    add_title_block(doc)
    add_callout(
        doc,
        "核心结论",
        "Cisco Skill Scanner 与 MCP Scanner 可以作为静态检测底座，但应通过独立适配器接入。平台自身必须掌握来源证明、统一证据模型、语义复核、失败闭锁和最终风险决策，不能把“扫描器没有发现”直接解释为“组件安全”。",
        fill=LIGHT_BLUE,
        accent=BLUE,
    )

    doc.add_heading("执行摘要", level=1)
    add_body(
        doc,
        "当前阶段建议优先交付一个可复现、可审计的静态供应链安全闭环：先固定输入制品并生成来源证明，再分别调用 Cisco Skill Scanner 和 Cisco MCP Scanner，随后把两套结果归一化为统一 Finding IR，通过规则证据与语义复核形成 ALLOW、REVIEW、BLOCK、UNKNOWN 四态决策。跨组件证据图和动态沙箱应在这一统一证据底座之上增量建设。",
    )
    add_body(
        doc,
        "现有复现结果表明，两个项目的工程框架、离线规则、报告输出和依赖漏洞检测均可用；但 Skill Scanner 在当前9个对抗样例上的召回率仅为28.57%，MCP Scanner 的6/6正确结果来自小规模基础样例，均不能据此声称已达到生产准确率。下一阶段的价值应集中在来源可信、证据融合、误报消歧和失败闭锁，而不是简单并排调用两个扫描器。",
    )
    add_section_map(doc, bullet_num)

    doc.add_heading("1. 项目背景与阶段目标", level=1)
    add_body(
        doc,
        "面向政企场景的 Agent 供应链不仅包含传统软件依赖，还包含 Skill 指令文件、脚本、二进制资产、MCP Tool 描述、Prompt、Resource、服务器配置和远程服务。攻击者可以在安装前、分发过程或运行时植入恶意逻辑，因此平台需要回答三个基础问题：组件从哪里来、组件包含什么、组件声称的能力与实际证据是否一致。",
    )
    add_body(
        doc,
        "本阶段目标：建立“可追溯输入—静态规则—语义复核—风险门禁—审计输出”的最小可用闭环。", bold_prefix="本阶段目标："
    )
    for text in [
        "完成 Skill、MCP 对象和依赖包的统一采集与只读留存。",
        "复用 Cisco 扫描器形成离线静态规则能力。",
        "建立受约束的 LLM 语义复核，不让模型独立作出阻断决定。",
        "建立来源证明和扫描可复现证明。",
        "输出统一 Finding IR、可解释决策和可回放审计记录。",
    ]:
        add_numbered_paragraph(doc, text, bullet_num)

    doc.add_heading("2. 范围、非目标与证据边界", level=1)
    add_table(
        doc,
        ["类别", "本阶段纳入", "本阶段不承诺"],
        [
            ["静态检测", "Skill 文件、脚本、文档、字节码、命令链；MCP Tool/Prompt/Resource；Python 依赖漏洞", "证明所有恶意逻辑均可被规则覆盖"],
            ["语义检测", "对静态证据进行上下文消歧、意图判断和描述—行为一致性复核", "让 LLM 单独决定安全或替代确定性规则"],
            ["来源证明", "仓库、提交、标签、哈希、签名状态、锁文件与扫描器版本", "仅凭哈希证明发布者可信"],
            ["跨组件", "预留统一节点、边和证据字段", "本阶段完成完整跨 Skill/MCP 攻击链推理"],
            ["动态验证", "定义接口、事件模型和后续验收条件", "本阶段完成高强度系统调用级沙箱"],
        ],
        [1500, 3930, 3930],
    )
    add_callout(
        doc,
        "边界原则",
        "凡是扫描器异常、输出为空、证据不足或来源不可确认的对象，状态必须是 UNKNOWN 或 REVIEW，而不是 SAFE。",
        fill="FFF4F2",
        accent=RED,
    )

    doc.add_heading("3. Cisco 基线选型与复用边界", level=1)
    doc.add_heading("3.1 Skill Scanner", level=2)
    add_body(
        doc,
        "Skill Scanner 适合承担 Skill 文件包的静态安检。建议复用 SkillLoader、StaticAnalyzer、BytecodeAnalyzer、PipelineAnalyzer、Policy 与 JSON/SARIF Reporter。其 BehavioralAnalyzer 可作为补充静态证据，但名称中的“行为分析”主要指 AST、控制流和污点传播，不代表真正执行 Skill。",
    )
    doc.add_heading("3.2 MCP Scanner", level=2)
    add_body(
        doc,
        "MCP Scanner 适合承担 MCP Server 对象采集、Tool/Prompt/Resource 离线快照扫描、YARA 规则、生产就绪性检查和依赖漏洞检测。远程与 stdio 扫描应首先产出不可变快照，再进入离线扫描，避免远程服务变化导致结果不可复现。",
    )
    add_table(
        doc,
        ["能力", "Skill Scanner", "MCP Scanner", "平台侧责任"],
        [
            ["文件/文档规则", "核心能力", "有限", "统一严重度与证据格式"],
            ["Tool/Prompt/Resource", "不面向 MCP 协议", "核心能力", "采集快照、内容哈希"],
            ["依赖漏洞", "非主链路", "pip-audit", "失败闭锁、漏洞归并"],
            ["语义分析", "可选 LLM/Meta", "可选 LLM/Meta", "隔离提示、证据约束、复核策略"],
            ["来源证明", "部分仓库获取信息", "部分包下载校验", "完整 provenance 与签名验证"],
            ["跨组件/动态", "启发式或预留", "局部源码入口", "证据图、组合风险与沙箱"],
        ],
        [1500, 2310, 2310, 3240],
    )

    doc.add_heading("4. 建议总体架构", level=1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    p.add_run().add_picture(str(architecture), width=Inches(6.35))
    cap = doc.add_paragraph(style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.add_run("图1  建议集成架构与阶段扩展关系")

    add_body(
        doc,
        "架构采用“采集与可信固化—检测执行器—统一证据—语义复核—策略决策”五层分离。两个 Cisco 扫描器作为可替换 Worker，主平台不依赖其内部 Finding 结构。这样既能固定当前版本，又能在后续升级上游或加入其他扫描器时保持接口稳定。",
    )
    for text in [
        "采集层：获取 Git 仓库、Skill 包、MCP 快照和依赖清单。",
        "来源证明层：计算哈希，记录提交、签名、锁文件和获取时间。",
        "检测执行层：分别调用 Skill Scanner 与 MCP Scanner 独立环境。",
        "证据归一化层：转换成统一 Finding IR，保留原始结果。",
        "语义与策略层：复核意图和误报，执行四态风险门禁。",
    ]:
        add_numbered_paragraph(doc, text, decimal_num)

    doc.add_heading("5. 静态规则扫描设计", level=1)
    doc.add_heading("5.1 Skill 扫描流水线", level=2)
    add_body(
        doc,
        "推荐顺序为：结构与 Manifest → 引用和文件清单 → YARA/文本规则 → 字节码一致性 → Shell/Pipeline → 可选静态数据流 → 统一去重和严重度策略。扫描路径必须是只读副本；压缩包解压需要文件数量、体积、层级和路径穿越限制。",
    )
    doc.add_heading("5.2 MCP 扫描流水线", level=2)
    add_body(
        doc,
        "推荐先通过 MCP 协议采集 tools/list、prompts/list、resources/list 或 resources/read，再把原始 JSON 与服务器标识共同固化。离线阶段扫描描述、参数 Schema、资源文本和服务器 Instructions，并对 requirements.txt/pyproject.toml 执行依赖漏洞审计。",
    )
    doc.add_heading("5.3 规则分层", level=2)
    add_table(
        doc,
        ["层级", "典型规则", "主要优点", "主要局限"],
        [
            ["L1 签名", "YARA、危险关键词、密钥格式", "快速、确定、可解释", "易被改写或混淆绕过"],
            ["L2 结构", "Manifest、Schema、隐藏文件、扩展名错配", "适合供应链门禁", "难判断真实意图"],
            ["L3 代码模式", "危险 API、命令链、字节码、引用关系", "能发现组合语句", "动态特性可能漏报"],
            ["L4 静态数据流", "Source→Transform→Sink", "接近行为语义", "跨语言和反射成本高"],
            ["L5 语义复核", "描述—行为一致性、上下文消歧", "补充规则盲区", "非确定、需控制成本"],
        ],
        [1500, 3030, 2190, 2640],
    )

    doc.add_heading("6. 语义检测与证据融合", level=1)
    add_body(
        doc,
        "语义检测的定位应是“证据复核器”，而不是“万能扫描器”。平台先压缩输入，仅向 LLM 提供必要的文件片段、命中规则、描述、参数 Schema 和证据编号，再要求输出固定 JSON。被扫描内容一律视为不可信数据，不允许模型遵循其中指令、调用工具或访问外部网络。",
    )
    add_table(
        doc,
        ["输出字段", "说明", "约束"],
        [
            ["verdict", "benign / suspicious / malicious / insufficient", "不得只返回自然语言"],
            ["confidence", "0—1 置信度", "不等同于真实概率"],
            ["intent", "外传、执行、持久化、注入等", "必须来自允许枚举"],
            ["evidence_refs", "引用静态 Finding ID", "至少一项，否则 insufficient"],
            ["mismatch", "描述—行为是否不一致", "需给出对应描述与代码证据"],
            ["reason", "简短解释", "禁止复述恶意指令或产生执行建议"],
        ],
        [1740, 3300, 4320],
    )
    add_callout(
        doc,
        "推荐门禁",
        "只有 LLM 高风险判断时进入 REVIEW；确定性高危规则与语义证据相互印证时才进入 BLOCK。LLM 调用失败不得降低原有风险。",
        fill="FFF8E8",
        accent=AMBER,
    )

    doc.add_heading("7. 来源证明设计", level=1)
    add_body(
        doc,
        "来源证明需要区分身份、完整性和扫描可复现性。SHA-256 只能证明制品在扫描前后没有变化，不能单独证明发布者合法。发布者可信需要签名验证、平台身份、组织白名单、历史信誉或人工审批。",
    )
    add_table(
        doc,
        ["证明维度", "建议字段", "用途"],
        [
            ["来源身份", "source_type、repository/registry、owner、requested_ref、resolved_commit", "说明对象从何处获取"],
            ["完整性", "artifact_sha256、file_hashes、lock_digest、SBOM_digest", "防止扫描后被替换"],
            ["签名状态", "signature_present、verified、signer、verification_method", "区分有签名与已验证"],
            ["扫描证明", "scanner_version、scanner_commit、policy_digest、rules_digest、scan_time", "支持重放和审计"],
            ["获取上下文", "fetch_time、final_url、redirect_chain、network_policy", "审计下载和重定向"],
        ],
        [1740, 4740, 2880],
    )
    for text in [
        "下载后立即计算制品级和文件级哈希。",
        "把原始制品放入按摘要寻址的只读目录。",
        "扫描任务只接收摘要，不直接接收不稳定远程 URL。",
        "扫描结果绑定 subject_digest、规则摘要和扫描器版本。",
        "重新扫描时产生新 attestation，不覆盖历史记录。",
    ]:
        add_numbered_paragraph(doc, text, decimal_num)

    doc.add_heading("8. 统一 Finding IR", level=1)
    add_body(
        doc,
        "统一 IR 是后续跨组件分析能否成功的关键。平台应保留原始 Cisco 结果，同时抽取稳定字段。严重度、置信度、扫描状态和来源可信度应是不同维度，避免用一个总分掩盖失败原因。",
    )
    add_table(
        doc,
        ["字段组", "核心字段", "说明"],
        [
            ["主体", "component_id、component_type、subject_digest", "唯一确定被扫描对象"],
            ["检测", "engine、analyzer、rule_id、category", "保留来源和规则身份"],
            ["风险", "severity、confidence、status", "status 必须支持 UNKNOWN/PARTIAL"],
            ["证据", "file、line、snippet、evidence_hash", "用于解释和复核"],
            ["来源", "repository、commit、signature、artifact_sha256", "与 provenance 记录关联"],
            ["语义", "verdict、intent、mismatch、evidence_refs", "只能引用已有证据"],
            ["审计", "scanner_version、policy_digest、started_at、finished_at", "支持复现实验"],
        ],
        [1500, 4050, 3810],
    )
    add_body(
        doc,
        "推荐对象类型：Skill、Script、Asset、MCPServer、Tool、Prompt、Resource、Package、SensitiveData、NetworkEndpoint。推荐关系类型：CONTAINS、REFERENCES、IMPORTS、CALLS、READS、WRITES、EXECUTES、DOWNLOADS、SENDS_TO。这些字段即使第一阶段暂时不用，也应在 Schema 中预留。",
    )

    doc.add_heading("9. 风险决策与失败闭锁", level=1)
    add_table(
        doc,
        ["状态", "进入条件", "平台动作"],
        [
            ["ALLOW", "来源满足策略；所有必需扫描器成功；没有超阈值发现", "允许进入测试或受限部署"],
            ["REVIEW", "只有语义风险；来源弱；中危证据冲突；需要人工确认", "进入审核队列"],
            ["BLOCK", "确定性高危规则；凭据外传；下载执行；恶意依赖；多证据一致", "阻断安装或接入"],
            ["UNKNOWN", "扫描器失败、超时、结果为空、输出不合法、关键证据缺失", "默认不放行，允许重试"],
        ],
        [1380, 4740, 3240],
    )
    add_body(
        doc,
        "必须单独保存 analyzer_status。某个分析器失败时，其他分析器可以继续运行，但总任务只能是 PARTIAL 或 UNKNOWN，不能因为 findings 为空而自动标为 SAFE。命令返回码、日志错误、JSON 是否存在、JSON 是否非空以及预期分析器是否真正执行都要进入完整性校验。",
    )

    doc.add_heading("10. 工程集成方案", level=1)
    add_body(
        doc,
        "两个项目不建议安装到同一个 Python 环境。当前复现中 Skill Scanner 使用 Python 3.11，MCP Scanner 使用 Python 3.13，并存在 LiteLLM、Rust 构建和依赖版本差异。推荐通过独立 Worker、子进程或容器调用，以 JSON 作为边界。",
    )
    add_table(
        doc,
        ["模块", "建议职责", "输入/输出"],
        [
            ["orchestrator", "任务编排、超时、重试、状态机", "ArtifactRef → ScanJob"],
            ["provenance", "获取、哈希、签名、SBOM、只读固化", "Source → Attestation"],
            ["skill_adapter", "调用 Skill Scanner、校验输出、归一化", "SkillPath → Finding[]"],
            ["mcp_adapter", "采集 MCP 快照、静态扫描、依赖审计", "Snapshot → Finding[]"],
            ["semantic_worker", "受约束语义复核", "EvidenceBundle → SemanticVerdict"],
            ["policy_engine", "四态门禁与例外策略", "Findings + Provenance → Decision"],
            ["audit_store", "保存原始制品、原始结果和归一化记录", "不可变审计链"],
        ],
        [1740, 4080, 3540],
    )
    add_body(
        doc,
        "调用子进程时使用参数数组而不是拼接 Shell 字符串；限制 CPU、内存、时间和输出大小；扫描目录默认只读；API Key 只注入到需要的 Worker；日志中禁止输出密钥和完整敏感片段。",
    )

    doc.add_heading("11. 当前复现证据与结论", level=1)
    add_table(
        doc,
        ["验证项", "结果", "解释边界"],
        [
            ["Skill 官方核心测试", "139 passed，6 skipped，1 xfailed", "证明目标核心路径在当前环境可运行"],
            ["MCP 官方核心测试", "114 passed", "覆盖 YARA、规则和依赖分析器"],
            ["Skill 对抗样例", "TP=2，TN=1，FP=1，FN=5；Recall=28.57%", "说明默认离线召回不足"],
            ["MCP 静态基础样例", "3良性+3恶意，6/6正确", "样本很小，不代表生产准确率"],
            ["依赖漏洞样例", "urllib3 1.24.1 检出14条；修复样例0条", "漏洞库会随时间变化"],
            ["云端/LLM能力", "未验证", "缺少 Cisco、LLM、VirusTotal Key"],
        ],
        [2040, 3000, 4320],
    )
    add_callout(
        doc,
        "阶段结论",
        "两个 Cisco 项目达到“可作为底座”的标准，但未达到“可直接作为最终检测器”的标准。平台价值应体现在来源可信、证据归一化、语义消歧和闭锁决策。",
        fill=LIGHT_BLUE,
        accent=BLUE,
    )

    doc.add_heading("12. 评测与验收设计", level=1)
    add_body(
        doc,
        "评测数据应按攻击机制分层，而不是只统计总准确率。至少包含良性业务 Skill、带攻击示例的安全文档、下载执行、凭据外传、描述—行为不一致、混淆载荷、Prompt Injection、恶意依赖和扫描失败样例。所有恶意样例应使用无效域名、模拟密钥和静态扫描，避免测试数据本身造成伤害。",
    )
    add_table(
        doc,
        ["指标", "计算/检查内容", "阶段建议目标"],
        [
            ["Recall", "恶意样例中被 BLOCK/REVIEW 覆盖的比例", "先达到 ≥80%，再降低 REVIEW"],
            ["Precision", "被判风险样例中真实恶意比例", "≥75%，重点分析文档语境误报"],
            ["F1", "Precision 与 Recall 的调和平均", "持续高于单一 Cisco 离线基线"],
            ["Fail-closed", "扫描失败是否全部进入 UNKNOWN", "100%"],
            ["Provenance coverage", "是否具有摘要、来源和扫描版本", "必需字段100%"],
            ["Reproducibility", "相同制品与规则是否得到一致静态结果", "确定性模块100%"],
            ["Latency", "单个中型 Skill/MCP 快照耗时", "在样例集上建立P50/P95基线"],
        ],
        [1680, 4260, 3420],
    )
    add_body(
        doc,
        "以上数值是下一阶段工程目标，不是当前已经达到的结果。评测报告应同时列出混淆矩阵、失败任务数、各攻击类别召回率和每个分析器的增量贡献。",
    )

    doc.add_heading("13. 分阶段开发计划", level=1)
    add_table(
        doc,
        ["阶段", "核心交付", "验收证据"],
        [
            ["第1周：统一入口", "制品模型、来源记录、摘要寻址存储、两个 Adapter", "相同制品可重放；原始结果完整保存"],
            ["第2周：静态闭环", "规则扫描、依赖审计、统一 Finding、四态门禁", "失败样例全部 UNKNOWN；基础报告可生成"],
            ["第3周：语义复核", "EvidenceBundle、固定 JSON Schema、Prompt 隔离、缓存", "安全文档误报下降；语义结果可追溯"],
            ["第4周：评测展示", "分层样例、指标脚本、策略配置、演示界面和视频", "一键复现；混淆矩阵和审计链完整"],
            ["后续：组合攻击", "证据图、跨 Skill/MCP 关系、风险路径", "检出单组件低危但组合高危的链"],
            ["可选：动态验证", "沙箱、模拟网络、文件/进程/环境变量事件", "静态推断与动态事件可关联"],
        ],
        [1740, 4440, 3180],
    )

    doc.add_heading("14. 主要风险与应对", level=1)
    add_table(
        doc,
        ["风险", "表现", "应对"],
        [
            ["规则绕过", "同义改写、混淆、动态导入", "规则+语义；保留 UNKNOWN；持续回归样例"],
            ["LLM不稳定", "输出漂移、成本、被注入", "固定Schema、无工具、温度控制、缓存、多证据门禁"],
            ["来源伪造", "URL可信但内容被替换", "摘要寻址、签名验证、获取后只读固化"],
            ["扫描器失败开放", "错误被当作SAFE", "输出完整性校验、必需分析器清单、UNKNOWN"],
            ["依赖冲突", "两个项目无法共用环境", "独立Worker/容器、固定提交和锁文件"],
            ["范围失控", "过早投入沙箱和跨语言动态分析", "先交付静态闭环，再按评测缺口扩展"],
        ],
        [1680, 3600, 4080],
    )

    doc.add_heading("15. 下一阶段开发清单", level=1)
    checklist = [
        "定义 ArtifactRef、ProvenanceRecord、Finding、ScanJob、Decision 五个核心 Schema。",
        "完成 Skill Scanner Adapter：版本自检、超时、JSON 校验、Finding 归一化。",
        "完成 MCP Snapshot Collector 与 MCP Scanner Adapter。",
        "实现按 SHA-256 寻址的只读制品目录。",
        "记录 scanner_commit、policy_digest、rules_digest 和 subject_digest。",
        "建立 COMPLETED_SAFE、COMPLETED_UNSAFE、PARTIAL、FAILED_UNKNOWN 状态机。",
        "实现 LLM EvidenceBundle 和固定输出 Schema，不允许 Tool Calling。",
        "建立至少30—50个分层样例，包含良性语境和失败样例。",
        "输出混淆矩阵、分类别召回、失败闭锁率和耗时。",
        "为跨组件证据图预留节点、边、证据引用和时间信息。",
    ]
    for item in checklist:
        add_numbered_paragraph(doc, item, bullet_num)

    doc.add_page_break()
    doc.add_heading("附录A：建议统一结果样例", level=1)
    sample = """{
  "finding_id": "F-001",
  "component_id": "sha256:...",
  "component_type": "skill",
  "engine": "cisco-skill-scanner",
  "analyzer": "pipeline_analyzer",
  "rule_id": "SENSITIVE_DATA_EXFILTRATION",
  "category": "data_exfiltration",
  "severity": "HIGH",
  "confidence": 0.92,
  "status": "confirmed",
  "evidence": {
    "file": "scripts/upload.py",
    "line": 18,
    "snippet_hash": "sha256:..."
  },
  "provenance_ref": "P-001",
  "scanner": {
    "version": "2.0.13.dev3",
    "commit": "4dee903...",
    "policy_digest": "sha256:..."
  }
}"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.right_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.0
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHT_GRAY)
    p_pr.append(shd)
    r = p.add_run(sample)
    set_run_font(r, latin="Consolas", east_asia="Microsoft YaHei", size=8.5)

    doc.add_heading("附录B：建议任务接口", level=1)
    add_table(
        doc,
        ["接口", "输入", "输出"],
        [
            ["POST /artifacts", "来源URL/本地包/注册表坐标", "ArtifactRef + ProvenanceRecord"],
            ["POST /scans", "artifact_id、scanner_profile、policy_id", "ScanJob"],
            ["GET /scans/{id}", "任务ID", "状态、原始结果、统一Finding"],
            ["POST /semantic-review", "EvidenceBundle", "SemanticVerdict"],
            ["POST /decisions", "Finding[]、Provenance、Policy", "ALLOW/REVIEW/BLOCK/UNKNOWN"],
        ],
        [2520, 3660, 3180],
    )

    doc.add_heading("附录C：参考来源与本地证据", level=1)
    sources = [
        ("Cisco Skill Scanner 官方仓库", "https://github.com/cisco-ai-defense/skill-scanner"),
        ("Cisco MCP Scanner 官方仓库", "https://github.com/cisco-ai-defense/mcp-scanner"),
        ("赛题材料", "XA-202620 面向政企场景的大模型智能体安全关键技术研究(6).pdf"),
        ("调研材料", "主要厂商供应链与Agent安全技术原理解读报告.pdf"),
        ("论文精读", "Agent_Skill与MCP安全_8篇必读论文精读报告_2026.docx"),
        ("论文清单", "Agent_Skill与MCP安全核心论文清单_2026.docx"),
        ("本地复现报告", str(ROOT / "REPRODUCTION_REPORT.md")),
        ("机器可读可用性结论", str(ROOT / "results" / "availability_summary.json")),
    ]
    for label, target in sources:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r1 = p.add_run(f"{label}：")
        set_run_font(r1, bold=True)
        if target.startswith("http"):
            add_hyperlink(p, target, target)
        else:
            r2 = p.add_run(target)
            set_run_font(r2, color=MUTED)

    doc.add_heading("附录D：版本说明", level=1)
    add_body(
        doc,
        "本报告以2026-07-31完成的本地复现为证据基础：Skill Scanner 固定提交4dee90371890ff23e1b21ea974e02847eacaa464；MCP Scanner 固定提交51966cce214ae057e69c3a672307911f5026e255。后续升级上游版本时，应重新运行官方核心测试、对抗样例和输出完整性检查。",
    )

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    result = build_document()
    print(result)
