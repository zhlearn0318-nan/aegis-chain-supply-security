"""Build a Chinese technical report about Cisco's two agent security scanners."""
from __future__ import annotations

from datetime import date
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from build_next_stage_report import (
    AMBER, BLUE, CALLOUT, GREEN, LIGHT_BLUE, MUTED, NAVY, RED,
    add_body, add_callout, add_hyperlink, add_numbered_paragraph,
    add_numbering_definitions, add_table, configure_section, configure_styles,
    set_run_font,
)

ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs"
ASSET_DIR = OUT_DIR / "_report_assets"
OUTPUT = OUT_DIR / "Cisco_Skill_Scanner与MCP_Scanner功能架构及集成角色分析报告.docx"
ARCHITECTURE = ASSET_DIR / "target_architecture.png"


def title_block(doc):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(4)
    set_run_font(p.add_run("技术调研与集成设计报告"), size=11, bold=True, color=BLUE)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(6)
    set_run_font(p.add_run("Cisco Skill Scanner 与 MCP Scanner"), size=25, bold=True, color=NAVY)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(18)
    set_run_font(p.add_run("功能架构、能力边界及在 Agent 供应链安全平台中的角色"), size=14, color=MUTED)
    for label, value in [
        ("报告用途", "下一阶段系统设计、模块拆分与开发验收参考"),
        ("分析对象", "Cisco Skill Scanner；Cisco MCP Scanner"),
        ("建议范围", "静态规则扫描、语义复核、来源证明、统一风险门禁"),
        ("版本日期", str(date.today())),
    ]:
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(3)
        set_run_font(p.add_run(f"{label}："), bold=True, color=NAVY)
        set_run_font(p.add_run(value))
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    add_callout(doc, "核心结论", "两个项目可以作为你的静态检测底座，但应被封装为底层检测执行器。来源证明、统一 Finding、四态风险决策、跨组件攻击图和最终准入策略应由你的平台自研。", fill=LIGHT_BLUE, accent=BLUE)


def link_source(doc, label, url):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    set_run_font(p.add_run(f"{label}："), bold=True)
    add_hyperlink(p, "官方 GitHub 仓库", url)


def intro(doc, bullet):
    doc.add_heading("1. 报告目的与阅读结论", 1)
    add_body(doc, "本报告面向一个月独立开发周期，解释 Cisco Skill Scanner 和 Cisco MCP Scanner 的扫描对象、内部架构、主要功能与工程边界，并给出二者在 Agent 供应链安全平台中的推荐角色。")
    for text in [
        "Skill Scanner 负责 Skill 制品、说明文档、脚本和资源的静态安全检测。",
        "MCP Scanner 负责 MCP 工具、提示词、资源、服务器说明和依赖包的安全检测。",
        "两者输出的是风险证据，不是安全证书；未发现风险不等于目标一定安全。",
        "你的系统负责可信采集、来源证明、证据归一化、风险门禁和跨组件分析。",
    ]: add_numbered_paragraph(doc, text, bullet)
    doc.add_heading("2. 两类扫描对象的区别", 1)
    add_table(doc, ["比较维度", "Agent Skill", "MCP Server"], [
        ["本质", "交给智能体的一组说明、脚本、配置和资源", "向智能体暴露工具、提示词和资源的服务"],
        ["主要入口", "SKILL.md、Python/Shell/JS、配置、二进制", "Tools、Prompts、Resources、Server Instructions"],
        ["典型风险", "隐藏指令、恶意脚本、凭据读取、下载执行、数据外传", "工具描述投毒、提示注入、危险能力、依赖漏洞、动态响应"],
        ["供应链关注点", "仓库、提交、作者、文件完整性、制品内容", "服务发布者、配置、依赖、连接地址、运行时行为"],
    ], [1800, 3780, 3780])


def skill_section(doc, bullet):
    doc.add_heading("3. Cisco Skill Scanner", 1)
    link_source(doc, "项目地址", "https://github.com/cisco-ai-defense/skill-scanner")
    add_body(doc, "Skill Scanner 是面向 AI Agent Skills 的多引擎安全扫描器。它把 Skill 视为完整制品，不只检查某个 Python 文件，还会加载 Skill 结构、读取说明文档、脚本、配置和资源，再把它们交给不同分析器。")
    doc.add_heading("3.1 逻辑架构", 2)
    add_table(doc, ["层次", "主要组件", "作用"], [
        ["输入与加载", "loader、repo_fetcher、strict_structure", "读取本地目录或 GitHub 仓库，识别 Skill 结构与可分析文件"],
        ["扫描编排", "scanner、analyzer_factory、scan_policy", "选择分析器、执行扫描、应用策略并聚合结果"],
        ["检测引擎", "static、bytecode、pipeline、behavioral、LLM 等", "分别从规则、代码、数据流和语义角度产生风险发现"],
        ["结果输出", "JSON、SARIF、HTML、Markdown、Table", "服务于平台集成、人工审查和 CI/CD 门禁"],
    ], [1700, 3100, 4560])
    add_callout(doc, "默认核心", "不配置外部密钥时，最适合第一阶段使用的是 Static、Bytecode 和 Pipeline 三类离线分析。", fill=CALLOUT, accent=GREEN)
    doc.add_heading("3.2 主要功能", 2)
    add_table(doc, ["分析器", "检查内容", "对你的价值"], [
        ["Static", "YAML、YARA、Python 规则；危险命令、隐藏字符、外传、提示注入等", "第一阶段核心规则引擎，可扩展政企规则"],
        ["Bytecode", ".pyc 完整性、无源码字节码和难以审计的编译产物", "把不可审计制品识别为供应链风险"],
        ["Pipeline", "Shell 管道中的来源、传播与危险落点", "发现敏感文件经命令管道发送到网络等链路"],
        ["Behavioral", "控制流、调用关系、跨文件和污点传播", "作为后续静态数据流分析基础，不等于动态沙箱"],
        ["LLM", "说明和代码的语义风险、意图、描述与行为一致性", "作为规则结果的语义复核，不宜单独阻断"],
        ["Meta", "多分析器结果聚合、去重和误报复核", "可借鉴证据融合思路，但平台仍需统一决策层"],
        ["Trigger / Cross-skill", "触发描述模糊、多个 Skill 名称或触发范围重叠", "为后续跨 Skill 冲突分析提供种子能力"],
        ["VirusTotal / AI Defense", "二进制恶意软件记录、云端文本检测", "可选增强；需要密钥、联网和数据合规评估"],
    ], [1600, 4310, 3450], font_size=8.8)
    doc.add_heading("3.3 在你的系统中的推荐角色", 2)
    add_callout(doc, "推荐角色", "Skill 制品静态安全扫描工作节点。", fill=LIGHT_BLUE, accent=BLUE)
    for text in [
        "接收经过哈希固定的 Skill 快照，而不是直接扫描一个可能变化的远程地址。",
        "执行结构检查、规则扫描、字节码检查、命令管道分析和可选语义复核。",
        "完整保留原始扫描结果、扫描器版本、启用的分析器和规则版本。",
        "通过适配器把结果转换成平台统一 Finding，而不是把原始 is_safe 直接作为准入决定。",
    ]: add_numbered_paragraph(doc, text, bullet)
    doc.add_heading("3.4 不应承担的职责", 2)
    for text in [
        "不负责证明仓库发布者一定可信，也不自动验证 Git 提交签名。",
        "不负责建立完整 SBOM 和制品流转证明。",
        "不负责 Skill 与 MCP、数据库、网络出口之间的完整组合攻击判断。",
        "静态 Behavioral 分析不能替代隔离环境中的真实运行验证。",
    ]: add_numbered_paragraph(doc, text, bullet)


def mcp_section(doc, bullet):
    doc.add_page_break(); doc.add_heading("4. Cisco MCP Scanner", 1)
    link_source(doc, "项目地址", "https://github.com/cisco-ai-defense/mcp-scanner")
    add_body(doc, "MCP Scanner 面向 Model Context Protocol 服务及其能力面。它既可以读取离线 JSON，也可以连接 stdio 或远程 MCP Server，获取工具、提示词、资源和服务器说明，再调用规则、语义、依赖漏洞和就绪性分析器。")
    doc.add_heading("4.1 逻辑架构", 2)
    add_table(doc, ["层次", "主要组件", "作用"], [
        ["输入适配", "本地配置、离线 JSON、stdio、远程 URL", "发现或连接 MCP Server，获取 Tools、Prompts、Resources"],
        ["扫描编排", "Scanner、Config、MCP models", "把不同输入转换为统一对象并调度分析器"],
        ["检测引擎", "YARA、LLM、API、依赖漏洞、Readiness、Behavioral", "从内容、依赖、配置和代码行为多个角度生成风险证据"],
        ["服务接口", "CLI、Python SDK、REST API", "便于独立运行或被你的平台调用"],
    ], [1700, 3100, 4560])
    doc.add_heading("4.2 扫描入口", 2)
    add_table(doc, ["入口", "特点", "阶段建议"], [
        ["离线 JSON", "无需启动目标服务器，可重复、便于审计", "第一阶段优先"],
        ["已知客户端配置", "发现 Cursor、Claude、VS Code 等客户端配置的 MCP", "后续用于资产发现"],
        ["stdio Server", "由扫描器启动本地进程并读取能力", "放入受限环境后启用"],
        ["远程 Server", "连接 URL 并读取服务器实际暴露内容", "需要网络隔离、认证和超时控制"],
        ["项目或依赖文件", "扫描 requirements 或 Python 项目依赖", "第一阶段启用"],
    ], [1800, 4100, 3460])
    doc.add_heading("4.3 主要功能", 2)
    add_table(doc, ["分析能力", "检查内容", "对你的价值"], [
        ["YARA", "提示注入、工具投毒、命令/脚本/SQL 注入、凭据窃取、数据外传", "MCP 静态内容扫描核心"],
        ["对象扫描", "Tools、Prompts、Resources、Server Instructions", "覆盖攻击可能出现的所有 MCP 文本入口"],
        ["Vulnerable Package", "通过 pip-audit 检查 CVE、GHSA、PYSEC 和修复版本", "直接形成依赖供应链漏洞证据"],
        ["Readiness", "超时、重试、错误处理、错误 Schema、暴露能力数量", "作为配置安全和生产就绪性检查"],
        ["LLM", "描述投毒、隐式指令、能力与描述不一致", "作为语义复核和规则补充"],
        ["Behavioral", "Python/JavaScript 控制流、调用图、污点和跨文件分析", "为后续代码行为与数据流分析提供参考"],
        ["VirusTotal / API", "文件恶意软件记录和 Cisco 云端检测", "可选增强，需评估密钥、联网和代码外发"],
    ], [1900, 4140, 3320], font_size=8.8)
    doc.add_heading("4.4 在你的系统中的推荐角色", 2)
    add_callout(doc, "推荐角色", "MCP 接口描述、内容、配置和依赖包的安全扫描工作节点。", fill=LIGHT_BLUE, accent=BLUE)
    for text in [
        "第一阶段通过离线 JSON 扫描 Tools、Prompts、Resources 和服务器说明。",
        "通过 pip-audit 扫描依赖漏洞，并记录外部工具的退出码和执行状态。",
        "把 Readiness 结果归入配置安全类别，不与恶意行为结论混为一谈。",
        "后续再在隔离环境中连接 stdio 或远程 MCP Server，采集运行时暴露面。",
    ]: add_numbered_paragraph(doc, text, bullet)
    doc.add_heading("4.5 不应承担的职责", 2)
    for text in [
        "不负责验证 MCP 发布者身份、仓库签名和服务端实际部署版本。",
        "不负责统一管理 Skill 与 MCP 的权限关系和攻击路径。",
        "不应在外部依赖执行失败时仍输出允许结论。",
        "连接 MCP Server 不等于安全沙箱，恶意服务仍可能攻击扫描进程。",
    ]: add_numbered_paragraph(doc, text, bullet)


def platform_section(doc, bullet):
    doc.add_heading("5. 两个项目的互补关系", 1)
    add_table(doc, ["维度", "Skill Scanner", "MCP Scanner", "平台统一处理"], [
        ["扫描对象", "Skill 文档、代码和资源", "MCP 工具、提示词、资源、服务和依赖", "统一 ArtifactRef"],
        ["规则能力", "YAML、YARA、Python 规则", "以 YARA 和专项分析器为主", "统一规则编号和分类"],
        ["语义分析", "分析 Skill 说明和代码", "分析 MCP 描述和内容", "统一语义复核服务"],
        ["数据流", "Skill 代码和 Shell 管道", "Python/JS MCP 代码", "后续建立跨组件图"],
        ["结果表达", "项目自身 Finding/报告", "项目自身 Result/报告", "统一 Finding IR"],
        ["最终决策", "内部安全标记或阈值", "内部安全标记或状态", "ALLOW/REVIEW/BLOCK/UNKNOWN"],
    ], [1350, 2450, 2450, 3110], font_size=8.6)
    doc.add_heading("6. 在总体设计中的位置", 1)
    add_body(doc, "建议把 Cisco 项目放在检测执行层。输入制品先经过可信采集和来源证明，再由两个独立工作节点扫描，最后进入统一证据、语义复核和风险门禁。")
    if ARCHITECTURE.exists():
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.keep_with_next = True
        p.add_run().add_picture(str(ARCHITECTURE), width=Inches(6.35))
        cap = doc.add_paragraph("图1  建议系统架构与两个 Cisco 扫描器的位置", style="Caption"); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_callout(doc, "架构原则", "检测器负责产生证据，平台负责判断证据是否完整、可信以及是否足以阻断。", fill=CALLOUT, accent=GREEN)
    doc.add_heading("7. 你的平台必须自研的上层能力", 1)
    add_table(doc, ["模块", "需要实现的核心内容", "原因"], [
        ["可信采集", "固定仓库、提交、下载时间和制品摘要；保存只读快照", "防止扫描对象在扫描前后被替换"],
        ["来源证明", "发布者、Git 签名、文件哈希、锁文件、SBOM、扫描声明", "两个扫描器不能证明来源真实性"],
        ["统一 Finding IR", "统一类别、严重度、位置、证据、置信度、状态和版本", "屏蔽两个项目不同的数据格式"],
        ["策略门禁", "ALLOW、REVIEW、BLOCK、UNKNOWN 四态决策", "避免把扫描失败误判为安全"],
        ["语义复核", "受约束提示词、结构化输出、多次一致性和证据引用", "控制 LLM 不稳定性和误报"],
        ["跨组件图", "连接 Skill、MCP、数据、凭据、权限和网络出口", "发现单组件扫描无法识别的组合攻击"],
        ["动态验证", "隔离运行、网络控制、系统调用和数据流追踪", "验证静态推断是否会在真实环境发生"],
    ], [1600, 4620, 3140], font_size=8.8)


def implementation(doc, bullet, decimal):
    doc.add_heading("8. 第一阶段集成建议", 1)
    add_body(doc, "一个月独立开发时，目标应是形成可重复、可解释、失败闭锁的离线扫描闭环。")
    for step in [
        "定义 ArtifactRef、ProvenanceRecord、Finding、ScanJob 和 Decision 五个统一数据模型。",
        "建立 Skill 与 MCP 两个独立扫描工作节点，避免 Python 版本和依赖互相污染。",
        "Skill 节点启用 Static、Bytecode、Pipeline；MCP 节点启用离线 YARA、依赖漏洞和 Readiness。",
        "保存原始扫描 JSON、标准输出、错误输出、退出码、扫描器版本和规则版本。",
        "实现适配器，把两个项目的结果转换为统一 Finding IR。",
        "实现四态策略门禁，并用安全、恶意、异常三类样本验证。",
    ]: add_numbered_paragraph(doc, step, decimal)
    doc.add_heading("8.1 建议启用矩阵", 2)
    add_table(doc, ["能力", "第一阶段", "后续阶段", "说明"], [
        ["Skill 静态/字节码/管道", "启用", "持续扩展", "离线、确定性较高"],
        ["MCP 离线 YARA", "启用", "持续扩展", "优先扫描导出的对象 JSON"],
        ["依赖漏洞扫描", "启用", "加入 SBOM 关联", "必须验证 pip-audit 执行状态"],
        ["LLM 语义分析", "受限试验", "正式复核层", "不直接输出最终阻断"],
        ["远程/stdio MCP", "暂缓", "隔离后启用", "存在启动进程和网络交互风险"],
        ["VirusTotal/云端 API", "可不启用", "按合规启用", "注意代码、文件和数据外发"],
        ["动态沙箱", "不作为核心", "时间充足再做", "需要独立隔离和观测体系"],
    ], [2050, 1350, 1600, 4360], font_size=8.8)
    doc.add_heading("9. 关键工程约束", 1)
    add_callout(doc, "必须失败闭锁", "扫描器异常、外部工具缺失、连接超时、文件读取失败或结果不完整时，状态必须是 UNKNOWN，不能因为没有 Finding 就进入 ALLOW。", fill="FDECEC", accent=RED)
    for text in [
        "不要直接使用 result.is_safe 作为最终准入条件。",
        "每个分析器必须返回 executed、success、error 和 coverage 等执行元数据。",
        "Skill Scanner 与 MCP Scanner 使用独立运行环境；当前复现分别采用 Python 3.11 和 Python 3.13。",
        "扫描任务只接收固定摘要对应的本地只读快照，禁止在判定过程中自动更新目标仓库。",
        "原始结果不可覆盖，重新扫描产生新的 attestation，并绑定扫描器提交和规则摘要。",
    ]: add_numbered_paragraph(doc, text, bullet)
    doc.add_heading("10. 当前复现证据", 1)
    add_table(doc, ["验证项", "结果", "如何解释"], [
        ["Skill Scanner 核心测试", "139 passed，6 skipped，1 xfailed", "核心离线代码在当前环境可运行"],
        ["MCP Scanner 定向测试", "114 passed", "选定核心模块在当前环境可运行"],
        ["Skill 9 个夹具", "TP=2、TN=1、FP=1、FN=5", "小样本召回率较低，不能单独作为安全门禁"],
        ["MCP 6 个基础夹具", "6/6 判断正确", "仅说明基础样本可用，样本量不足以证明泛化能力"],
        ["依赖漏洞样本", "urllib3 1.24.1 检出14项；升级后为0", "pip-audit 链路可用，但必须检查执行失败"],
        ["云端与 LLM", "未验证", "缺少 Cisco、LLM 和 VirusTotal 密钥，不应写成已完成能力"],
    ], [2300, 2550, 4510], font_size=8.8)
    doc.add_heading("11. 后续演进路线", 1)
    add_table(doc, ["阶段", "建设重点", "Cisco 项目的作用"], [
        ["阶段一", "静态规则、来源证明、统一 Finding、风险门禁", "作为主要离线检测执行器"],
        ["阶段二", "语义复核、误报分析、证据融合", "提供 LLM/Meta 能力与原始证据"],
        ["阶段三", "Skill-MCP-数据-权限攻击图", "提供节点风险和局部数据流信息"],
        ["阶段四", "隔离沙箱、网络控制、系统调用和动态数据流", "作为动态验证前的静态筛选器"],
    ], [1500, 4450, 3410])
    doc.add_heading("12. 最终结论", 1)
    add_body(doc, "两个 Cisco 项目适合被集成，但不能直接等同于完整供应链安全平台。Skill Scanner 解决 Skill 制品内容检测，MCP Scanner 解决 MCP 能力面和依赖检测；你的核心创新应位于可信采集、来源证明、统一证据模型、失败闭锁门禁和跨组件攻击链分析。")
    add_callout(doc, "一句话定位", "Skill Scanner 是 Skill 安检机，MCP Scanner 是 MCP 安检机；你的平台是身份核验、证据中心、风险决策和攻击链研判系统。", fill="FFF8E8", accent=AMBER)


def appendix(doc):
    doc.add_heading("附录A：建议统一 Finding 字段", 1)
    add_table(doc, ["字段", "含义"], [
        ["finding_id", "平台内唯一风险发现编号"],
        ["artifact_id / subject_digest", "被扫描制品及其不可变摘要"],
        ["component_type", "skill、mcp_tool、mcp_prompt、mcp_resource、dependency 等"],
        ["scanner / analyzer", "产生结果的扫描器和具体分析器"],
        ["category / severity / confidence", "统一风险分类、严重度和置信度"],
        ["location / evidence", "文件、行号、对象名称、规则命中和证据片段"],
        ["execution_status", "是否执行、是否成功、错误信息和覆盖情况"],
        ["scanner_version / rule_digest", "扫描器版本与规则集合摘要"],
        ["decision", "ALLOW、REVIEW、BLOCK 或 UNKNOWN"],
    ], [2800, 6560])
    doc.add_heading("附录B：来源与版本", 1)
    add_body(doc, "本报告以本地复现的官方仓库代码、README、测试结果和可用性记录为依据。")
    link_source(doc, "Cisco Skill Scanner", "https://github.com/cisco-ai-defense/skill-scanner")
    add_body(doc, "本地复现提交：4dee90371890ff23e1b21ea974e02847eacaa464。")
    link_source(doc, "Cisco MCP Scanner", "https://github.com/cisco-ai-defense/mcp-scanner")
    add_body(doc, "本地复现提交：51966cce214ae057e69c3a672307911f5026e255。")


def build():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document(); configure_styles(doc); configure_section(doc.sections[0])
    doc.core_properties.title = "Cisco Skill Scanner 与 MCP Scanner 功能架构及集成角色分析报告"
    doc.core_properties.subject = "Agent 供应链安全平台下一阶段开发参考"
    doc.core_properties.author = "技术调研组"
    bullet, decimal = add_numbering_definitions(doc)
    title_block(doc); intro(doc, bullet); skill_section(doc, bullet); mcp_section(doc, bullet)
    platform_section(doc, bullet); implementation(doc, bullet, decimal); appendix(doc)
    doc.save(OUTPUT); return OUTPUT


if __name__ == "__main__": print(build())
